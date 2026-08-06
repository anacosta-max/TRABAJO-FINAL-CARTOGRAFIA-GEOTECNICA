# -*- coding: utf-8 -*-
"""
Genera el artículo en formato DYNA (revista de la Facultad de Minas,
Universidad Nacional de Colombia): papel carta, 2 columnas en el cuerpo,
Times New Roman, encabezados numerados, citación [n], tablas con título
arriba, figuras con título abajo.

Estructura siguiendo la guía del profesor (separar Resultados de
Discusión): Introducción, Área de estudio, Datos y métodos, Resultados,
Discusión, Conclusiones, Referencias.

Contenido: caracterización de modelos y evaluación de la susceptibilidad
a movimientos en masa en la cuenca de la quebrada El Circio (Abejorral,
Antioquia), con los 3 métodos desarrollados en este proyecto: AHP +
Combinado (heurístico), Peso de la Evidencia/WoE (bivariado) y Regresión
Logística (multivariado).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, "figuras")
OUT = os.path.join(BASE, "Articulo_Susceptibilidad_Quebrada_ElCircio.docx")

F16 = r"C:\Users\Angela Acosta\OneDrive\Documentos\TRABAJO_FINAL_CARTOGRAFIA\MODELOS\16_Validacion_ROC_Comparativa\figuras"
F06 = r"C:\Users\Angela Acosta\OneDrive\Documentos\TRABAJO_FINAL_CARTOGRAFIA\MODELOS\Resultados\06_Variables_Condicionantes\figuras"
F09 = r"C:\Users\Angela Acosta\OneDrive\Documentos\TRABAJO_FINAL_CARTOGRAFIA\MODELOS\Resultados\09_Exploracion_Seleccion_Variables\figuras_estratificado"
GH_MAPS = r"C:\Users\Angela Acosta\OneDrive\Documentos\GitHub\TRABAJO-FINAL-CARTOGRAFIA-GEOTECNICA\maps"
DATOS = r"C:\Users\Angela Acosta\OneDrive\Documentos\TRABAJO_FINAL_CARTOGRAFIA\DATOS PROCESADOS"

# ===========================================================================
# CONFIGURACIÓN DEL DOCUMENTO
# ===========================================================================
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10)
rpr = normal.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "Times New Roman")

MARGEN = Cm(1.69)
sec0 = doc.sections[0]
sec0.page_height = Cm(27.94)
sec0.page_width = Cm(21.59)
sec0.top_margin = MARGEN; sec0.bottom_margin = MARGEN
sec0.left_margin = MARGEN; sec0.right_margin = MARGEN


def set_columnas(section, num):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), "425")
    if qn("w:equalWidth") in cols.attrib:
        del cols.attrib[qn("w:equalWidth")]


def nueva_seccion(num_columnas):
    s = doc.add_section(WD_SECTION.CONTINUOUS)
    s.top_margin = MARGEN; s.bottom_margin = MARGEN
    s.left_margin = MARGEN; s.right_margin = MARGEN
    s.page_height = Cm(27.94); s.page_width = Cm(21.59)
    set_columnas(s, num_columnas)
    return s


set_columnas(sec0, 1)

contador_fig = [0]
contador_tab = [0]


def run_tnr(p, texto, negrita=False, italica=False, tam=10, color=None):
    r = p.add_run(texto)
    r.font.name = "Times New Roman"
    r.font.size = Pt(tam)
    r.bold = negrita
    r.italic = italica
    if color:
        r.font.color.rgb = color
    return r


def parrafo(texto, tam=10, sangria=True, justificar=True, espacio_after=0):
    p = doc.add_paragraph()
    if justificar:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if sangria:
        p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(espacio_after)
    run_tnr(p, texto, tam=tam)
    return p


def encabezado1(numero, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run_tnr(p, f"{numero} {texto}", negrita=True, tam=10)
    return p


def encabezado2(numero, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run_tnr(p, f"{numero} {texto}", negrita=True, italica=True, tam=10)
    return p


def set_cell_bg(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def figura(ruta, texto_caption, ancho_in=3.35, fuente="Elaboración propia"):
    contador_fig[0] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(); run.add_picture(ruta, width=Inches(ancho_in))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(1)
    run_tnr(cap, f"Fig. {contador_fig[0]}. {texto_caption}", tam=8)
    capf = doc.add_paragraph(); capf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    capf.paragraph_format.space_after = Pt(8)
    run_tnr(capf, f"Fuente: {fuente}", tam=8, italica=True)
    return contador_fig[0]


def tabla_num(headers, filas, texto_caption, tam=7.5, fuente="Elaboración propia"):
    contador_tab[0] += 1
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6); cap.paragraph_format.space_after = Pt(1)
    run_tnr(cap, f"Tabla {contador_tab[0]}. {texto_caption}", tam=8)

    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, hcell in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tnr(p, str(hcell), negrita=True, tam=tam)
    for fila in filas:
        cells = t.add_row().cells
        for i, val in enumerate(fila):
            cells[i].text = ""
            p = cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_tnr(p, str(val), tam=tam)
    fnt = doc.add_paragraph()
    fnt.paragraph_format.space_after = Pt(8)
    run_tnr(fnt, f"Fuente: {fuente}", tam=8, italica=True)
    return contador_tab[0]


def cita_bloque(items):
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.space_after = Pt(4)
        run_tnr(p, it, tam=9)


# ===========================================================================
# PORTADA (1 columna)
# ===========================================================================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tnr(t, "Characterization of Models and Evaluation of the Landslide Susceptibility "
           "in the El Circio Creek Basin, Abejorral, Antioquia",
        negrita=True, tam=13)
t.paragraph_format.space_after = Pt(2)

t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tnr(t2, "Caracterización de Modelos y Evaluación de la Susceptibilidad de Movimientos "
            "en Masa en la Cuenca de la Quebrada El Circio, Abejorral, Antioquia",
        negrita=True, tam=13)
t2.paragraph_format.space_after = Pt(10)

aut = doc.add_paragraph(); aut.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tnr(aut, "Angela Yineth Acosta Alfonso", tam=11)
aut.paragraph_format.space_after = Pt(2)

afil = doc.add_paragraph(); afil.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tnr(afil, "Programa de Ingeniería Geológica, Departamento de Geociencias y Medio "
              "Ambiente, Facultad de Minas, Universidad Nacional de Colombia, "
              "Medellín, Colombia.", tam=9)
afil.paragraph_format.space_after = Pt(14)

h = doc.add_paragraph(); run_tnr(h, "Abstract", negrita=True, tam=10)
h.paragraph_format.space_after = Pt(2)
parrafo(
    "This study characterizes three landslide susceptibility models in the El Circio creek "
    "basin (Abejorral, Antioquia, Colombia), 41.3 km2, using ten conditioning variables. A "
    "heuristic model (Analytic Hierarchy Process, AHP, combined with class weights derived "
    "from Frequency Ratio), a bivariate statistical model (Weight of Evidence, WoE), and a "
    "multivariate statistical model (logistic regression with stepwise variable selection) "
    "were built and validated against a 256-point inventory (128 landslides, 128 stable "
    "points) using ROC curves, AUC, and the DeLong test. Logistic regression achieved the "
    "highest discrimination (AUC=0.869), followed by WoE (AUC=0.819) and AHP+Combined "
    "(AUC=0.667); differences among the three models are statistically significant. "
    "Susceptibility maps were classified into three levels following the Colombian "
    "Geological Survey criterion (High=75% of landslides, Medium=+23%, Low=remainder).",
    sangria=False, espacio_after=6,
)
kw = doc.add_paragraph()
run_tnr(kw, "Keywords: ", negrita=True, tam=10)
run_tnr(kw, "Landslide susceptibility; AHP; Weight of Evidence; Logistic regression; "
            "ROC curve; DeLong test; El Circio creek basin.", tam=10)
kw.paragraph_format.space_after = Pt(10)

h = doc.add_paragraph(); run_tnr(h, "Resumen", negrita=True, tam=10)
h.paragraph_format.space_after = Pt(2)
parrafo(
    "Este estudio caracteriza tres modelos de susceptibilidad a movimientos en masa en la "
    "cuenca de la quebrada El Circio (Abejorral, Antioquia, Colombia), de 41.3 km2, a "
    "partir de diez variables condicionantes. Se construyó y validó un modelo heurístico "
    "(Proceso Analítico Jerárquico, AHP, combinado con pesos de clase por Frequency Ratio), "
    "uno estadístico bivariado (Peso de la Evidencia, WoE) y uno estadístico multivariado "
    "(regresión logística con selección de variables por pasos), contra un inventario de "
    "256 puntos (128 movimientos, 128 estables), mediante curvas ROC, AUC y la prueba de "
    "DeLong. La regresión logística obtuvo la mayor discriminación (AUC=0.869), seguida por "
    "WoE (AUC=0.819) y AHP+Combinado (AUC=0.667); las diferencias entre los tres modelos "
    "son estadísticamente significativas. Los mapas de susceptibilidad se clasificaron en "
    "tres niveles según el criterio del Servicio Geológico Colombiano (Alta=75% de los "
    "movimientos, Media=+23%, Baja=resto).",
    sangria=False, espacio_after=6,
)
kw2 = doc.add_paragraph()
run_tnr(kw2, "Palabras clave: ", negrita=True, tam=10)
run_tnr(kw2, "Susceptibilidad a movimientos en masa; AHP; Peso de la Evidencia; Regresión "
             "logística; Curva ROC; Prueba de DeLong; Cuenca quebrada El Circio.", tam=10)
kw2.paragraph_format.space_after = Pt(6)

fecha = doc.add_paragraph()
run_tnr(fecha, "© Angela Yineth Acosta Alfonso; Universidad Nacional de Colombia. "
               "Agosto de 2026. Medellín.", tam=8, italica=True)

# ===========================================================================
# CUERPO DEL ARTÍCULO (2 columnas)
# ===========================================================================
nueva_seccion(2)

# ---------------------------------------------------------------------
# 1. Introducción
# ---------------------------------------------------------------------
encabezado1("1", "Introducción")
parrafo(
    "Los movimientos en masa son uno de los procesos morfodinámicos más recurrentes en "
    "las zonas de ladera de los Andes colombianos, condicionados por la combinación de "
    "factores topográficos, geológicos, de cobertura y uso del suelo. La cuenca de la "
    "quebrada El Circio, en el municipio de Abejorral (Antioquia), presenta un relieve "
    "montañoso con pendientes altas que la hacen susceptible a este tipo de procesos, "
    "por lo que la zonificación de su susceptibilidad es una herramienta útil para la "
    "gestión del riesgo a escala municipal."
)
parrafo(
    "Este trabajo caracteriza y compara tres métodos de modelamiento de la susceptibilidad "
    "a movimientos en masa con fundamentos epistemológicos distintos: un método heurístico "
    "(Proceso Analítico Jerárquico, AHP, combinado con pesos de clase estadísticos), un "
    "método estadístico bivariado (Peso de la Evidencia, WoE) y un método estadístico "
    "multivariado (regresión logística). El objetivo es evaluar objetivamente el "
    "desempeño de cada modelo mediante curvas ROC, área bajo la curva (AUC) y la prueba "
    "estadística de DeLong, y generar los respectivos mapas de zonificación de la "
    "susceptibilidad siguiendo el criterio de clasificación del Servicio Geológico "
    "Colombiano (SGC) [7]."
)

# ---------------------------------------------------------------------
# 2. Área de estudio
# ---------------------------------------------------------------------
encabezado1("2", "Área de estudio")
parrafo(
    "La cuenca de la quebrada El Circio se localiza en el municipio de Abejorral, "
    "departamento de Antioquia, Colombia. Para la ubicación cartográfica inicial se "
    "utilizó como referencia la Plancha 167 del Instituto Geográfico Agustín Codazzi "
    "(IGAC). La delimitación de la cuenca se realizó en ambiente SIG a partir de un "
    "Modelo Digital de Elevación (DEM) ALOS PALSAR de 12.5 m de resolución espacial, "
    "mediante el procesamiento hidrológico estándar (corrección de depresiones, "
    "dirección de flujo, acumulación de flujo, definición del punto de cierre y "
    "generación de la cuenca con la herramienta Watershed)."
)
figura(f"{DATOS}/Mapa_Localizacion_Ver3.jpg",
       "Mapa de localización de la cuenca de la quebrada El Circio, "
       "municipio de Abejorral, Antioquia.", ancho_in=3.35,
       fuente="Elaboración propia, IGAC, Esri, HERE, Garmin, USGS, METI/NASA")
tabla_num(
    ["Parámetro", "Valor"],
    [
        ["Nombre de la cuenca", "Quebrada El Circio"],
        ["Municipio", "Abejorral, Antioquia"],
        ["Área", "41.3 km² (4 100 ha)"],
        ["Perímetro", "30.5 km"],
        ["Escala de trabajo", "1:25 000"],
        ["Modelo de elevación", "ALOS PALSAR, 12.5 m"],
        ["Cartografía base", "Plancha 167 IGAC"],
        ["Elevación (mín. – máx. – media)", "775 – 2 596 – 1 754 msnm"],
        ["Pendiente media", "26.9°"],
    ],
    "Parámetros generales de la cuenca de la quebrada El Circio.",
)
figura(f"{FIG}/dem.png", "Modelo de Elevación Digital (DEM) de la cuenca de la "
       "quebrada El Circio.", ancho_in=2.6)

parrafo(
    "La geología y la geomorfología del área, al igual que la cobertura y el uso "
    "actual del suelo, se obtuvieron del Plan de Ordenación y Manejo de la Cuenca "
    "(POMCA) del río Arma, recortadas al polígono de la cuenca de la quebrada El "
    "Circio. En la cuenca se identifican seis unidades geológicas, dominadas en área "
    "por el Neis Intrusivo Abejorral (40.4%) y el Complejo Cajamarca (esquistos y "
    "filitas, en conjunto 42.8% del área). Al cruzar el inventario con estas "
    "unidades se observa un patrón inverso al del área ocupada: el Complejo "
    "Cajamarca — Esquistos verdes concentra la mayor densidad de movimientos (0.94 "
    "movimientos/100 ha; 84% de los 50 puntos muestreados sobre esta unidad "
    "corresponden a movimiento), seguido de las Filitas y esquistos "
    "cuarzo-sericíticos del mismo complejo (0.80 movimientos/100 ha; 55% de 69 "
    "puntos), mientras que el Neis Intrusivo Abejorral, pese a ser la unidad de "
    "mayor área, presenta la menor densidad relativa entre las unidades con muestra "
    "suficiente (0.30 movimientos/100 ha; 31% de 96 puntos)."
)
tabla_num(
    ["Código", "Unidad geológica", "% del área"],
    [
        [1, "Neis Intrusivo de Pantanillo", "5.2%"],
        [2, "Complejo Cajamarca – Esquistos verdes", "13.1%"],
        [3, "Neis Intrusivo Abejorral", "40.4%"],
        [4, "Formación Abejorral", "10.8%"],
        [5, "Complejo Quebradagrande – Miembro Volcánico", "0.8%"],
        [6, "Complejo Cajamarca – Filitas y esquistos cuarzo-sericíticos", "29.7%"],
    ],
    "Unidades geológicas presentes en la cuenca y su proporción de área.",
)
figura(f"{FIG}/geologia.png", "Mapa geológico de la cuenca de la quebrada El Circio.",
       ancho_in=2.3, fuente="POMCA río Arma, elaboración propia")

# ---------------------------------------------------------------------
# 3. Datos y métodos
# ---------------------------------------------------------------------
encabezado1("3", "Datos y métodos")

encabezado2("3.1", "Inventario de movimientos en masa")
parrafo(
    "El inventario consolidado de movimientos en masa de la cuenca está compuesto por "
    "128 registros: 92 corresponden a información secundaria extraída del Plan de "
    "Ordenación y Manejo de la Cuenca (POMCA) del río Arma, y 36 fueron identificados "
    "mediante fotointerpretación propia de imágenes satelitales y fotografías aéreas "
    "multitemporales disponibles en Google Earth Pro. Del total, 126 corresponden a "
    "deslizamientos y 2 a flujos; 22 se encuentran activos y 106 inactivos (Tabla 3, "
    "Fig. 3-5)."
)
tabla_num(
    ["Variable", "Resultado"],
    [
        ["Total de movimientos inventariados", 128],
        ["De fotointerpretación propia", 36],
        ["De información secundaria (POMCA río Arma)", 92],
        ["Deslizamientos", 126],
        ["Flujos", 2],
        ["Movimientos activos", 22],
        ["Movimientos inactivos", 106],
    ],
    "Resumen estadístico del inventario de movimientos en masa.",
)
figura(f"{GH_MAPS}/Mapa_Tipo.jpg", "Mapa de tipo de movimiento en masa de la cuenca "
       "de la quebrada El Circio.", ancho_in=3.3)
figura(f"{GH_MAPS}/Mapa_Actividad.jpg", "Mapa de actividad de los movimientos en "
       "masa de la cuenca.", ancho_in=3.3)
figura(f"{GH_MAPS}/Mapa_Multitemporal.jpg", "Clasificación multitemporal de los "
       "movimientos en masa identificados por fotointerpretación.", ancho_in=3.3)

parrafo(
    "Para la variable dependiente del modelo se generaron puntos representativos de la "
    "zona de inicio de cada movimiento (vértice de mayor elevación de cada polígono, "
    "extraída mediante el DEM), a los cuales se les asignó el valor Y=1. Para "
    "representar la ausencia de movimiento (Y=0), se generaron inicialmente 128 puntos "
    "aleatorios dentro de la cuenca, excluyendo los polígonos de movimiento y un buffer "
    "de seguridad de 25 m (~2 celdas del DEM). Este conjunto inicial de puntos estables "
    "fue posteriormente refinado mediante un criterio de disimilitud ambiental "
    "(distancia en el espacio de las 10 variables condicionantes respecto a los puntos "
    "de movimiento) y estratificación espacial en una cuadrícula de 18 celdas, con el "
    "fin de evitar la concentración espacial de los puntos estables y mejorar su "
    "representatividad como contraste de los movimientos reales. La capa final "
    "(Puntos_Modelo_Balanceado) contiene 256 puntos: 128 con ocurrencia de movimiento y "
    "128 sin evidencia de ocurrencia, en proporción 1:1."
)

encabezado2("3.2", "Variables condicionantes")
parrafo(
    "Se evaluaron diez variables condicionantes: seis continuas (pendiente, aspecto, "
    "curvatura, flujo acumulado, elevación, distancia a drenajes) y cuatro categóricas "
    "(cobertura, uso actual del suelo, geología, geomorfología), todas con una "
    "resolución espacial de 12.5 m. Antes de construir los modelos se evaluó la "
    "capacidad individual de cada variable para discriminar entre celdas con y sin "
    "movimiento mediante pruebas t / Mann-Whitney (continuas) y chi-cuadrado "
    "(categóricas), con un nivel de significancia α=0.05, y se realizó un análisis de "
    "componentes principales (PCA) para evaluar la redundancia entre variables."
)
tabla_num(
    ["Variable", "Tipo", "Significativa (p<0.05)"],
    [
        ["Pendiente", "Continua", "Sí (p≈0.03)"],
        ["Aspecto", "Continua (circular)", "Sí (p<0.001)"],
        ["Curvatura", "Continua", "No"],
        ["Flujo acumulado", "Continua", "No"],
        ["Elevación", "Continua", "Sí (p<0.01)"],
        ["Distancia a drenajes", "Continua", "No"],
        ["Cobertura", "Categórica", "Sí (p<0.001)"],
        ["Uso actual del suelo", "Categórica", "Sí (p<0.001)"],
        ["Geología", "Categórica", "Sí (p<0.001)"],
        ["Geomorfología", "Categórica", "Sí (p<0.001)"],
    ],
    "Variables condicionantes evaluadas y su significancia estadística univariada.",
)
parrafo(
    "Las otras tres variables categóricas (cobertura, uso actual del suelo y "
    "geomorfología) provienen, al igual que la geología (Fig. 3), del POMCA del río "
    "Arma, recortado al polígono de la cuenca. Las cinco variables de terreno "
    "restantes derivadas del DEM (la elevación se presentó en la Fig. 2) se "
    "muestran en la Fig. 7."
)
figura(f"{FIG}/mosaico_continuas.png",
       "Variables continuas condicionantes derivadas del DEM: pendiente, aspecto, "
       "curvatura, flujo acumulado y distancia a drenajes.", ancho_in=3.35)

figura(f"{F09}/matriz_correlacion_pearson.png",
       "Matriz de correlación de Pearson entre las variables continuas.", ancho_in=3.3)
figura(f"{F06}/07_biplot_pc1_pc2.png", "Biplot del Análisis de Componentes Principales "
       "(PCA) de las variables condicionantes.", ancho_in=3.3)
interpretacion_placeholder = None  # (sin interpretación aquí; se reserva para Discusión)
parrafo(
    "Flujo acumulado, distancia a drenajes y curvatura no resultaron significativas en "
    "ningún diseño muestral evaluado, hallazgo consistente a lo largo del proyecto; el "
    "PCA confirmó que las 10 variables aportan información mayormente independiente, "
    "requiriéndose 7 de 10 componentes para explicar el 80% de la varianza."
)

encabezado2("3.3", "Modelo heurístico: AHP + Combinado")
parrafo(
    "El Proceso Analítico Jerárquico (AHP) [3] pondera cada variable (Wᵢ) a partir de una "
    "matriz de comparación por pares en la escala de Saaty, resuelta mediante el "
    "eigenvector principal, con verificación de consistencia (Razón de Consistencia, "
    "CR=0.020, consistente). El índice de susceptibilidad se calcula como la suma "
    "ponderada Sₙ=ΣWᵢ·w_cᵢ, donde w_c es el peso asignado a cada clase dentro de una "
    "variable (eq. 1)."
)
ec1 = doc.add_paragraph(); ec1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tnr(ec1, "Sₙ = Σ Wᵢ · w_cᵢ", italica=True, tam=10)
ec1.paragraph_format.space_after = Pt(6)
tabla_num(
    ["Variable", "Wᵢ (AHP)"],
    [
        ["Pendiente", "0.242"], ["Geología", "0.242"], ["Aspecto", "0.110"],
        ["Curvatura", "0.110"], ["Flujo acumulado", "0.110"], ["Elevación", "0.047"],
        ["Distancia a drenajes", "0.047"], ["Geomorfología", "0.047"],
        ["Cobertura", "0.022"], ["Uso actual del suelo", "0.022"],
    ],
    "Pesos de variable (Wᵢ) obtenidos con la matriz AHP.",
)
parrafo(
    "Para las cuatro variables categóricas, el peso de clase (w_c) se calculó con "
    "Frequency Ratio (FR), definido como w_n=L_r/A_r, donde L_r es el porcentaje de los "
    "128 movimientos reales contenido en la clase y A_r el porcentaje del área total de "
    "la cuenca que representa dicha clase, normalizado por el w_n máximo de la "
    "variable. Las seis variables continuas conservaron rangos de reclasificación de "
    "literatura especializada en la región andina colombiana."
)

encabezado2("3.4", "Modelo bivariado: Peso de la Evidencia (WoE)")
parrafo(
    "El método del Peso de la Evidencia [4], basado en el teorema de Bayes, calcula "
    "para cada clase de cada variable un peso positivo W+ (presencia de movimientos "
    "en la clase) y un peso negativo W- (ausencia), a partir de conteos de celdas con "
    "y sin movimiento sobre el ráster completo del dominio de estudio (eq. 2-3)."
)
ec2 = doc.add_paragraph(); ec2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tnr(ec2, "W⁺ = ln[Npix1/(Npix1+Npix3)] − ln[Npix2/(Npix2+Npix4)]", italica=True, tam=9)
ec2.paragraph_format.space_after = Pt(6)
parrafo(
    "El contraste C=W⁺−W⁻ resume la asociación neta de la clase con la ocurrencia de "
    "movimientos: valores positivos indican correlación directa y negativos, "
    "correlación inversa. El índice final de susceptibilidad corresponde a la suma de "
    "los pesos de todas las variables en cada celda."
)

encabezado2("3.5", "Modelo multivariado: Regresión Logística")
parrafo(
    "La regresión logística estima la probabilidad de ocurrencia de movimiento (Y=1) "
    "en función de un conjunto de variables predictoras, mediante máxima verosimilitud "
    "(eq. 4). La selección del subconjunto óptimo de variables se realizó por los tres "
    "métodos clásicos de selección por pasos (forward selection, backward elimination "
    "y stepwise bidireccional), usando el Criterio de Información de Akaike (AIC) como "
    "criterio de decisión, tratando cada variable categórica como un bloque completo."
)
ec3 = doc.add_paragraph(); ec3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_tnr(ec3, "P(Y=1) = 1 / (1 + e^−(β₀+Σβᵢxᵢ))", italica=True, tam=10)
ec3.paragraph_format.space_after = Pt(6)
parrafo(
    "Los tres métodos convergieron exactamente al mismo modelo de cinco variables "
    "(uso actual del suelo, geología, geomorfología, curvatura y distancia a "
    "drenajes), reduciendo el AIC de 356.9 (modelo nulo) a 268.9 y aumentando el "
    "pseudo R² de McFadden de 0 a 0.36. La convergencia idéntica de los tres métodos "
    "de búsqueda es evidencia de que el modelo resultante es un óptimo robusto por "
    "AIC. Se verificó además que las 10 variables no pueden incluirse simultáneamente: "
    "cobertura y uso actual del suelo están casi perfectamente anidadas en la muestra, "
    "lo que produce una matriz de diseño no identificable."
)

encabezado2("3.6", "Evaluación de modelos")
parrafo(
    "La capacidad de discriminación de cada modelo se evaluó mediante la curva ROC "
    "(Receiver Operating Characteristic), que grafica la sensibilidad (TPR) contra "
    "1-especificidad (FPR) para todos los umbrales posibles del índice de "
    "susceptibilidad, y el área bajo la curva (AUC), interpretada según la escala "
    "cualitativa de Swets [5]: AUC>0.9 alto, AUC>0.7 moderado, AUC>0.5 aceptable. Para "
    "los modelos heurístico y bivariado (AHP+Combinado, WoE), cuyos pesos no se "
    "ajustan por optimización contra los datos, se calculó una única curva ROC sobre "
    "el 100% del inventario. Para la regresión logística, que sí se ajusta a los "
    "datos y puede sobreajustarse, se aplicó una partición estratificada 80/20 "
    "(entrenamiento/validación), reajustando los coeficientes solo con el 80% y "
    "evaluando la curva ROC tanto en el propio conjunto de entrenamiento como en el "
    "20% nunca visto por el modelo."
)
parrafo(
    "La significancia estadística de las diferencias de AUC entre modelos se evaluó "
    "con la prueba de DeLong [6], que compara curvas ROC de forma pareada sobre el "
    "mismo conjunto de sujetos; para esta comparación se usó el modelo de regresión "
    "logística ajustado con el 100% de los datos, para igualar la base de comparación "
    "con AHP+Combinado y WoE."
)
parrafo(
    "Finalmente, los mapas de susceptibilidad se clasificaron en tres niveles "
    "siguiendo el criterio de tasa de éxito acumulada recomendado por el Servicio "
    "Geológico Colombiano [7]: se ordenan los 128 movimientos reales por su score de "
    "mayor a menor y se toma como umbral \"Alta\" el score que captura el 75% de los "
    "movimientos (TPR=0.75); el umbral \"Media\" corresponde al 23% siguiente "
    "(TPR acumulado=0.98); \"Baja\" es el resto."
)

# ---------------------------------------------------------------------
# 4. Resultados
# ---------------------------------------------------------------------
encabezado1("4", "Resultados")
parrafo(
    "La Tabla 5 y la Fig. 9 presentan la curva ROC y el AUC de los tres modelos "
    "evaluados sobre el mismo conjunto de 256 puntos del inventario, incluyendo el "
    "ajuste de la regresión logística con el 100% de los datos (n=256), usado como "
    "referencia para la prueba de DeLong de la Tabla 6."
)
tabla_num(
    ["Modelo", "Curva", "n", "AUC", "Swets (1988)"],
    [
        ["AHP + Combinado", "Única", 256, 0.667, "Aceptable"],
        ["WoE", "Única", 256, 0.819, "Moderado"],
        ["Regresión Log.", "Completo (100%, ref. DeLong)", 256, 0.869, "Moderado"],
        ["Regresión Log.", "Entrenamiento (80%)", 204, 0.898, "Moderado"],
        ["Regresión Log.", "Validación (20%)", 51, 0.729, "Moderado"],
    ],
    "AUC y clasificación cualitativa (Swets, 1988) de los tres modelos.",
)
figura(f"{F16}/roc_comparacion_3_modelos_delong.png",
       "Curvas ROC de los tres modelos sobre los mismos 256 puntos "
       "(AHP+Combinado, WoE y regresión logística ajustada con el 100% de los datos).",
       ancho_in=3.3)
figura(f"{F16}/roc_regresion_logistica_train_test.png",
       "Curvas ROC de la regresión logística con partición 80/20: "
       "entrenamiento (curva de predicción) y validación (curva de éxito).",
       ancho_in=3.3)
parrafo(
    "Según la escala cualitativa de Swets [5], el modelo AHP+Combinado (AUC=0.667) "
    "se clasifica como \"aceptable\" (mejor que el azar, pero por debajo del umbral "
    "moderado), mientras que WoE (AUC=0.819) y la regresión logística —tanto en su "
    "ajuste completo (AUC=0.869) como en entrenamiento (AUC=0.898) y validación "
    "(AUC=0.729)— se clasifican como \"moderados\" (moderadamente precisos); ninguno "
    "de los tres modelos alcanza el rango \"alto\" (AUC>0.9)."
)

parrafo("La Tabla 6 presenta los resultados de la prueba de DeLong para las tres comparaciones pareadas posibles.")
tabla_num(
    ["Comparación", "ΔAUC", "z", "p"],
    [
        ["AHP+Combinado vs. WoE", "−0.152", "−4.21", "<0.0001"],
        ["AHP+Combinado vs. Regresión", "−0.202", "−5.77", "<0.0001"],
        ["WoE vs. Regresión", "−0.050", "−2.75", "0.0059"],
    ],
    "Resultados de la prueba de DeLong (todas las diferencias son significativas, p<0.05).",
)

parrafo(
    "Como métrica complementaria al AUC, la Tabla 6b presenta la matriz de "
    "confusión de cada modelo evaluada en un umbral de referencia (la mediana del "
    "score para AHP+Combinado y WoE; 0.5 para la regresión logística, siguiendo el "
    "criterio del capítulo). La regresión logística alcanza la mayor exactitud "
    "tanto en entrenamiento (83.3%) como en validación (72.6%), seguida de WoE "
    "(76.6%); AHP+Combinado presenta la exactitud más baja (61.7%), consistente "
    "con su menor AUC."
)
tabla_num(
    ["Modelo", "Conjunto (n)", "Umbral", "TP", "FP", "TN", "FN", "Exactitud"],
    [
        ["AHP + Combinado", "Único (256)", "0.492 (mediana)", 79, 49, 79, 49, "61.7%"],
        ["WoE", "Único (256)", "−0.675 (mediana)", 98, 30, 98, 30, "76.6%"],
        ["Regresión Log.", "Entrenamiento (204)", "0.5", 82, 14, 88, 20, "83.3%"],
        ["Regresión Log.", "Validación (51)", "0.5", 18, 7, 19, 7, "72.6%"],
    ],
    "Matriz de confusión y exactitud de los tres modelos en su umbral de referencia.",
)

parrafo(
    "La clasificación en tres niveles (Alta, Media, Baja) siguiendo el criterio del "
    "SGC (Tabla 7, Fig. 11) muestra que, por construcción, la clase \"Alta\" de los "
    "tres mapas contiene el 75% de los movimientos reales; la diferencia entre "
    "modelos radica en cuánta área de la cuenca se requiere para lograrlo: "
    "AHP+Combinado necesita el 63.7% del área de la cuenca en la clase Alta (7.8% "
    "Baja, 28.5% Media) para capturar ese 75% de movimientos, frente a un 48.2% en "
    "regresión logística (12.5% Baja, 39.2% Media) y solo un 42.2% en WoE (3.3% "
    "Baja, 54.6% Media) — WoE es, por tanto, el modelo que logra la mayor "
    "concentración espacial de la susceptibilidad alta."
)
tabla_num(
    ["Modelo", "% área Baja", "% área Media", "% área Alta"],
    [
        ["AHP + Combinado", "7.8%", "28.5%", "63.7%"],
        ["WoE", "3.3%", "54.6%", "42.2%"],
        ["Regresión Logística", "12.5%", "39.2%", "48.2%"],
    ],
    "Distribución de área por clase de susceptibilidad (criterio SGC: 75/23/2).",
)
figura(f"{F16}/mapa_susceptibilidad_AHP_3clases.png",
       "Mapa de susceptibilidad clasificado en tres niveles (criterio SGC) — "
       "Modelo heurístico AHP + Combinado.", ancho_in=2.6)
figura(f"{F16}/mapa_susceptibilidad_WoE_3clases.png",
       "Mapa de susceptibilidad clasificado en tres niveles (criterio SGC) — "
       "Modelo bivariado Peso de la Evidencia (WoE).", ancho_in=2.6)
figura(f"{F16}/mapa_susceptibilidad_Regresion_3clases.png",
       "Mapa de susceptibilidad clasificado en tres niveles (criterio SGC) — "
       "Modelo multivariado Regresión Logística.", ancho_in=2.6)

# ---------------------------------------------------------------------
# 5. Discusión
# ---------------------------------------------------------------------
encabezado1("5", "Discusión")
parrafo(
    "La regresión logística obtuvo la mayor capacidad de discriminación (AUC=0.869 "
    "con el 100% de los datos; AUC=0.729 en validación con el 20% nunca visto), "
    "seguida por WoE (AUC=0.819) y, con una diferencia estadísticamente significativa "
    "respecto a los dos anteriores (prueba de DeLong, p<0.0001), por AHP+Combinado "
    "(AUC=0.667). La prueba de DeLong confirma que esta jerarquía es estadísticamente "
    "robusta y no un efecto del azar: las tres comparaciones pareadas resultan "
    "significativas (AHP+Combinado vs. WoE, z=-4.21, p<0.0001; AHP+Combinado vs. "
    "Regresión, z=-5.77, p<0.0001; y, con menor margen pero también significativa, "
    "WoE vs. Regresión, z=-2.75, p=0.0059), de modo que ningún par de modelos puede "
    "considerarse estadísticamente equivalente en su capacidad de discriminación. "
    "Esta jerarquía es además consistente con el grado en que cada método "
    "incorpora evidencia estadística local: la regresión logística ajusta sus "
    "coeficientes por máxima verosimilitud sobre el 100% de las variables "
    "seleccionadas; WoE calcula sus pesos directamente de conteos de movimientos del "
    "dominio completo; AHP+Combinado, en cambio, solo incorporó datos reales (vía "
    "Frequency Ratio) en las cuatro variables categóricas, mientras las seis "
    "continuas permanecieron ancladas a rangos de literatura genérica, lo que limita "
    "su capacidad de adaptarse a las particularidades de la cuenca de El Circio."
)
parrafo(
    "Un hallazgo que merece explicación es que la regresión logística, con el AUC más "
    "alto, requiere más área (48.2%) que WoE (42.2%) para capturar el mismo 75% de "
    "movimientos reales en la clase \"Alta\". El AUC mide qué tan bien un modelo "
    "ordena, en general, los puntos de movimiento por encima de los estables; no mide "
    "qué tan concentrada queda la probabilidad predicha sobre el mapa completo. Como "
    "tres de las cinco variables del modelo de regresión son categóricas (uso del "
    "suelo, geología, geomorfología), la probabilidad asignada tiende a ser casi "
    "uniforme dentro de polígonos completos de esas categorías, generando bloques "
    "extensos y contiguos de alta susceptibilidad. WoE, al combinar aditivamente diez "
    "variables (incluidas las seis continuas, con variación gradual pixel a pixel), "
    "produce un mapa más heterogéneo y, en este caso, más concentrado espacialmente."
)
parrafo(
    "El desempeño de AHP+Combinado mejoró sustancialmente al reemplazar los pesos de "
    "clase de literatura por Frequency Ratio en las variables categóricas: el AUC "
    "pasó de 0.542 (100% literatura) a 0.667, y la prueba t de separación de medias "
    "entre movimientos y puntos estables pasó de no significativa (p=0.283) a "
    "altamente significativa (p=1.6×10⁻⁶). Esto confirma que la principal limitación "
    "del método heurístico puro en esta cuenca no es su estructura (la combinación "
    "lineal ponderada), sino la falta de anclaje de sus pesos a evidencia local, y "
    "sugiere que extender el mismo criterio a las variables continuas podría reducir "
    "aún más la brecha con los métodos estadísticos."
)
parrafo(
    "Como limitación general, el AUC de entrenamiento y de validación de la regresión "
    "logística difieren en 0.169, indicativo de cierto sobreajuste esperable dado el "
    "tamaño reducido del conjunto de validación (51 puntos) frente al número de "
    "parámetros del modelo (21, la mayoría de variables categóricas). Asimismo, la "
    "evaluación de AHP+Combinado y WoE sobre el 100% del inventario, sin partición, "
    "es más comparable en espíritu a un AUC de entrenamiento que a uno de validación, "
    "por lo que la comparación más conservadora entre los tres modelos es la que usa "
    "el AUC de validación de la regresión (0.729), que de igual forma permanece por "
    "encima de WoE en términos de discriminación estadísticamente robusta según la "
    "prueba de DeLong."
)

# ---------------------------------------------------------------------
# 6. Conclusiones
# ---------------------------------------------------------------------
encabezado1("6", "Conclusiones")
parrafo(
    "Se caracterizaron y evaluaron tres modelos de susceptibilidad a movimientos en "
    "masa en la cuenca de la quebrada El Circio, con fundamentos metodológicos "
    "distintos: heurístico (AHP+Combinado), bivariado (WoE) y multivariado (regresión "
    "logística). La regresión logística presentó la mayor capacidad de discriminación "
    "(AUC=0.869), seguida de WoE (AUC=0.819) y de AHP+Combinado (AUC=0.667); las tres "
    "diferencias son estadísticamente significativas según la prueba de DeLong "
    "(p<0.05 en las tres comparaciones)."
)
parrafo(
    "El desempeño de un método heurístico depende críticamente de que sus pesos de "
    "clase incorporen evidencia estadística local: al recalcular los pesos de las "
    "variables categóricas de AHP+Combinado mediante Frequency Ratio, su AUC mejoró "
    "de 0.542 a 0.667. Un mayor AUC no garantiza, sin embargo, un mapa más "
    "concentrado espacialmente: el modelo de regresión logística, con el mejor AUC, "
    "requirió más área para capturar el 75% de los movimientos reales que el modelo "
    "WoE, debido al efecto de bloque de sus variables categóricas."
)
parrafo(
    "Se recomienda para trabajos futuros extender la calibración estadística de "
    "pesos (Frequency Ratio u otro criterio basado en datos) a las variables "
    "continuas del método heurístico, y ampliar el inventario de movimientos para "
    "reducir la incertidumbre de la partición de validación de la regresión "
    "logística."
)

# ---------------------------------------------------------------------
# Referencias
# ---------------------------------------------------------------------
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
run_tnr(p, "Referencias", negrita=True, tam=10)
cita_bloque([
    "[1] Aristizábal, E., Libro Guía del curso Cartografía Geotécnica, Universidad "
    "Nacional de Colombia, Departamento de Geociencias y Medio Ambiente, Facultad de "
    "Minas, Medellín, Colombia, 2022. Disponible en: "
    "https://edieraristizabal.github.io/Libro_cartoGeotecnia/",
    "[2] Varnes, D.J., and International Association of Engineering Geology, "
    "Landslide hazard zonation: a review of principles and practice, UNESCO, Paris, "
    "Francia, 1984. ISBN 9231018957.",
    "[3] Saaty, T.L., The Analytic Hierarchy Process: Planning, Priority Setting, "
    "Resource Allocation, McGraw-Hill, New York, USA, 1980.",
    "[4] Bonham-Carter, G.F., Geographic Information Systems for Geoscientists: "
    "Modelling with GIS, Pergamon Press, Oxford, UK, 1994.",
    "[5] Swets, J.A., Measuring the accuracy of diagnostic systems, Science, "
    "240(4857), pp. 1285-1293, 1988. DOI: https://doi.org/10.1126/science.3287615",
    "[6] DeLong, E.R., DeLong, D.M., and Clarke-Pearson, D.L., Comparing the areas "
    "under two or more correlated receiver operating characteristic curves: a "
    "nonparametric approach, Biometrics, 44(3), pp. 837-845, 1988. "
    "DOI: https://doi.org/10.2307/2531595",
    "[7] Servicio Geológico Colombiano (SGC), Guía metodológica para estudios de "
    "amenaza, vulnerabilidad y riesgo por movimientos en masa, SGC, Bogotá, "
    "Colombia, 2017.",
])

doc.save(OUT)
print(f"\nArtículo completo guardado: {OUT}")
