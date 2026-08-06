# -*- coding: utf-8 -*-
"""
Informe Word corto que interpreta los resultados de validación ROC/AUC de
los 4 métodos del proyecto (AHP+Combinado, Índices, WoE, Regresión
Logística), siguiendo la distinción de las notas de clase (EVALUACIÓN DE
MODELOS, 29/07/2026) y el capítulo oficial "Evaluación del modelo" del
libro:
  - Métodos heurísticos (AHP+Combinado, Índices, WoE): 1 sola curva ROC,
    sin partición.
  - Métodos estadísticos/de datos (Regresión Logística): 2 curvas,
    entrenamiento (80%, "curva de predicción") y validación
    (20%, "curva de éxito").

Fuente de AHP actualizada: proyecto_susceptibilidad_AHP (matriz AHP +
Frequency Ratio DESCARTADO, solo asignación directa de pesos de clase,
para calzar con el ejemplo de referencia del curso).
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
DIR_FIG = os.path.join(BASE, "figuras")
DIR_TAB = os.path.join(BASE, "tablas")
OUT = os.path.join(BASE, "Informe_Validacion_ROC.docx")

# ===========================================================================
# 0. CARGAR TABLAS
# ===========================================================================
conf_ahp = pd.read_csv(f"{DIR_TAB}/ahp+combinado_matriz_confusion_resumen.csv").iloc[0]
conf_woe = pd.read_csv(f"{DIR_TAB}/woe_matriz_confusion_resumen.csv").iloc[0]
auc_reg = pd.read_csv(f"{DIR_TAB}/regresion_auc_train_test.csv")
metricas_extra = pd.read_csv(f"{DIR_TAB}/metricas_adicionales_r_swets.csv")
delong = pd.read_csv(f"{DIR_TAB}/delong_comparacion_pareada.csv")

auc_ahp = conf_ahp["AUC"]
auc_woe = conf_woe["AUC"]
auc_train = auc_reg.loc[auc_reg["conjunto"].str.contains("Entrenamiento"), "AUC"].iloc[0]
auc_test = auc_reg.loc[auc_reg["conjunto"].str.contains("Prueba"), "AUC"].iloc[0]


def fila_metrica(modelo, contiene_conjunto=None):
    sub = metricas_extra[metricas_extra["modelo"] == modelo]
    if contiene_conjunto:
        sub = sub[sub["conjunto"].str.contains(contiene_conjunto)]
    return sub.iloc[0]


r_ahp = fila_metrica("AHP+Combinado")
r_woe = fila_metrica("WoE")
r_reg_train = fila_metrica("Regresión Logística", "Entrenamiento")
r_reg_test = fila_metrica("Regresión Logística", "Prueba")

# ===========================================================================
# 1. TABLA RESUMEN COMPARATIVA
# ===========================================================================
resumen = pd.DataFrame([
    {"modelo": "AHP + Combinado", "tipo": "Heurístico", "curva": "Única (sin partición)",
     "n_puntos": 256, "AUC": round(auc_ahp, 3), "clasificación (Swets 1988)": r_ahp["clasificacion_Swets_1988"]},
    {"modelo": "WoE (Peso de la Evidencia)", "tipo": "Heurístico", "curva": "Única (sin partición)",
     "n_puntos": 256, "AUC": round(auc_woe, 3), "clasificación (Swets 1988)": r_woe["clasificacion_Swets_1988"]},
    {"modelo": "Regresión Logística", "tipo": "Estadístico/datos", "curva": "Entrenamiento (80%, predicción)",
     "n_puntos": 204, "AUC": round(auc_train, 3), "clasificación (Swets 1988)": r_reg_train["clasificacion_Swets_1988"]},
    {"modelo": "Regresión Logística", "tipo": "Estadístico/datos", "curva": "Validación (20%, éxito)",
     "n_puntos": 51, "AUC": round(auc_test, 3), "clasificación (Swets 1988)": r_reg_test["clasificacion_Swets_1988"]},
])
resumen.to_csv(f"{DIR_TAB}/resumen_comparativo_AUC.csv", index=False, encoding="utf-8-sig")
print(resumen.to_string(index=False))

# ===========================================================================
# 2. INFORME WORD
# ===========================================================================
AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x40, 0x40, 0x40)
VERDE_OK = RGBColor(0x1E, 0x7B, 0x34)
ROJO_ALERTA = RGBColor(0xC0, 0x00, 0x00)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
for s in doc.sections:
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p


def para(text, bold=False, italic=False, size=11, color=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def interpretacion(text, alerta=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run("Alerta: " if alerta else "Interpretación: ")
    r.bold = True; r.italic = True; r.font.size = Pt(10.5)
    r.font.color.rgb = ROJO_ALERTA if alerta else VERDE_OK
    r2 = p.add_run(text)
    r2.italic = True; r2.font.size = Pt(10.5); r2.font.color.rgb = GRIS
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(11)


def set_cell_bg(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def tabla(headers, rows, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; r = p.add_run(str(hd))
        r.bold = True; r.font.size = Pt(font_size); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1F4E79")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]; r = p.add_run(str(val))
            r.font.size = Pt(font_size)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def tabla_desde_csv(ruta_csv, font_size=9, decimales=3):
    df = pd.read_csv(ruta_csv)
    df_fmt = df.copy()
    for col in df_fmt.columns:
        if pd.api.types.is_float_dtype(df_fmt[col]):
            df_fmt[col] = df_fmt[col].round(decimales)
    return tabla([str(c) for c in df_fmt.columns], df_fmt.values.tolist(), font_size=font_size)


def imagen(nombre, width_in=5.6, caption=None):
    ruta = os.path.join(DIR_FIG, nombre)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(ruta, width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption); cr.italic = True; cr.font.size = Pt(9.5); cr.font.color.rgb = GRIS
        cp.paragraph_format.space_after = Pt(4)


def salto():
    doc.add_page_break()


# --- Portada ---
doc.add_paragraph().paragraph_format.space_before = Pt(50)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Validación de Modelos: Curvas ROC y AUC")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = AZUL

st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = st.add_run("Comparación AHP+Combinado, Índices, WoE y Regresión Logística")
r2.font.size = Pt(14); r2.font.color.rgb = GRIS
st.paragraph_format.space_before = Pt(10)

st3 = doc.add_paragraph(); st3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = st3.add_run("Cuenca del río Arma (Abejorral, Antioquia)")
r4.font.size = Pt(12); r4.font.color.rgb = GRIS
st3.paragraph_format.space_before = Pt(30)
doc.add_page_break()

# ===========================================================================
# 1. Metodología de validación
# ===========================================================================
h("1. Criterio de validación aplicado", 1)
para(
    "Siguiendo las notas de clase (\"Evaluación de modelos\", 29/07/2026) y el capítulo "
    "oficial \"Evaluación del modelo\" del libro, la validación se hace distinto según "
    "el tipo de modelo:"
)
bullets([
    "Métodos heurísticos (AHP+Combinado, Índices, WoE — Peso de la Evidencia): sus "
    "pesos NO se ajustan por optimización iterativa contra los datos, así que no aplica "
    "un esquema de entrenamiento/validación. Basta con UNA curva ROC calculada sobre el "
    "inventario completo — así lo indican las notas: \"Para los métodos heurísticos (por "
    "ejemplo peso de la evidencia y asignación de pesos), se debe hacer al menos una "
    "curva ROC\".",
    "Regresión Logística (método estadístico/de datos): sí se ajusta (máxima "
    "verosimilitud) contra los datos, por lo que puede sobreajustarse. Se exige "
    "partición 80/20: \"para regresión logística, mínimo hacer la partición de 80-20\", "
    "confirmado también por el capítulo del libro como el estándar en estudios "
    "colombianos. Se reporta la curva de entrenamiento (\"curva de predicción\") y la de "
    "validación (\"curva de éxito\", sobre el 20% nunca visto por el modelo).",
])
para(
    "El capítulo del libro complementa lo anterior con TRES métricas de evaluación: la "
    "matriz de confusión, la curva ROC/AUROC, y la Distancia a la clasificación perfecta "
    "(r = √(FPR² + (1-TPR)²) — distancia euclidiana hasta la esquina superior izquierda "
    "ideal; menor r es mejor). También reporta la escala cualitativa de Swets (1988), "
    "usada en las guías SafeLand/JTC-1: AUC > 0.9 = altamente preciso, AUC > 0.7 = "
    "moderadamente preciso, AUC > 0.5 = mejor que el azar (aceptable), AUC ≤ 0.5 = no "
    "aceptable. Las tres métricas se calculan para los tres modelos comparados en este "
    "informe: AHP+Combinado, WoE y Regresión Logística (el método de Índices, casi "
    "redundante con AHP+Combinado, se excluye de esta comparación por decisión del "
    "usuario)."
)

# ===========================================================================
# 2. AHP + Combinado
# ===========================================================================
h("2. Método AHP + Combinado (heurístico)", 1)
para(
    "Fuente actualizada: proyecto_susceptibilidad_AHP. A diferencia de la versión "
    "anterior de este informe, esta vez el peso de las clases (w_c) se asigna "
    "ÚNICAMENTE por criterio experto — se descartó el uso de Frequency Ratio para "
    "calibrar los pesos, por decisión explícita de la autora, para ajustarse al ejemplo "
    "de referencia del curso. El índice de susceptibilidad (S_n = Σ W_i·w_c_i) se "
    "recalculó en Python punto por punto (sin arcpy) usando las mismas reglas de "
    "reclasificación de los CSV del proyecto; se verificó que el t-test de validación "
    "coincide con el ya reportado en el README del proyecto (t=1.077 vs. 1.076, "
    "p=0.2825 vs. 0.2828 — solo redondeo)."
)
imagen("roc_ahp.png", width_in=5.0, caption="Figura 2.1. Curva ROC — AHP + Combinado.")
interpretacion(
    f"AUC = {auc_ahp:.3f} → \"{r_ahp['clasificacion_Swets_1988']}\" según Swets (1988) — "
    "apenas por encima de 0.5, es decir, el modelo discrimina SOLO marginalmente mejor "
    "que una clasificación aleatoria. Esto es consistente con el t-test ya reportado en "
    "el README del proyecto (p=0.28, NO significativo): la media del índice en puntos "
    "de movimiento (0.548) y en puntos estables (0.536) es casi idéntica. Distancia a "
    f"la clasificación perfecta r = {r_ahp['distancia_clasificacion_perfecta_r']:.3f} "
    "(la más alta —peor— de los 3 modelos evaluados en este informe).",
    alerta=True,
)

# ===========================================================================
# 3. WoE
# ===========================================================================
h("3. Método WoE — Peso de la Evidencia (heurístico)", 1)
para(
    "Sin cambios respecto a la validación anterior. Se extrajo el valor del raster "
    "susceptibilidad_WoE.tif (14_WoE_Susceptibilidad, pesos W+ calculados de conteos de "
    "píxeles del dominio completo) en los mismos 256 puntos del inventario."
)
imagen("roc_woe.png", width_in=4.6, caption="Figura 3.1. Curva ROC — WoE.")
interpretacion(
    f"AUC = {auc_woe:.3f} → \"{r_woe['clasificacion_Swets_1988']}\" (Swets 1988) — muy "
    f"superior a AHP+Combinado ({auc_ahp:.3f}). Distancia a la clasificación perfecta "
    f"r = {r_woe['distancia_clasificacion_perfecta_r']:.3f}, muy por debajo (mejor) que "
    "la de AHP+Combinado. La diferencia clave: los pesos W+ de WoE se calculan "
    "matemáticamente de la evidencia estadística del inventario (una fórmula cerrada "
    "basada en conteos reales de movimientos por clase), mientras que AHP+Combinado, en "
    "esta versión, depende completamente del criterio experto sin ningún anclaje a los "
    "datos — la diferencia de AUC es la evidencia más clara de cuánto aporta dejar que "
    "los datos calibren los pesos."
)

# ===========================================================================
# 4. Regresión logística
# ===========================================================================
h("4. Regresión Logística (estadístico, con partición 80/20)", 1)
para(
    "Modelo final ya seleccionado por AIC en 13_Regresion_Logistica (uso_actual + "
    "geología + geomorfología + curvatura + dist_drenajes). Partición 80/20 "
    "estratificada (semilla=42); categorías raras agrupadas SOLO con los conteos del "
    "80% de entrenamiento; coeficientes reajustados solamente con ese 80%."
)
imagen("roc_regresion_logistica_train_test.png", width_in=5.0,
       caption="Figura 5.1. Curvas ROC — entrenamiento vs. validación.")
interpretacion(
    f"AUC entrenamiento = {auc_train:.3f} (\"{r_reg_train['clasificacion_Swets_1988']}\"), "
    f"AUC validación = {auc_test:.3f} (\"{r_reg_test['clasificacion_Swets_1988']}\") — "
    "diferencia de 0.169, indicando cierto sobreajuste (esperable con solo 51 puntos de "
    "prueba y 21 parámetros), pero incluso en datos nunca vistos el modelo sigue siendo "
    "moderadamente preciso. Se excluyó 1 punto del 20% de prueba por caer en una "
    "categoría de geología nunca observada en entrenamiento (limitación documentada, no "
    "se fuerza una predicción inválida)."
)
salto()

# ===========================================================================
# 5. Comparación estadística por pares (DeLong)
# ===========================================================================
h("5. Comparación estadística por pares — prueba de DeLong", 1)
para(
    "Responde a la Actividad 3 del capítulo del libro: \"traza las curvas ROC de los... "
    "métodos... en el mismo gráfico. ¿Cuál tiene mayor AUC? ¿Es la diferencia "
    "estadísticamente significativa (prueba DeLong)?\". Se comparan los 3 modelos sobre "
    "EXACTAMENTE los mismos 256 puntos (para regresión se usa el modelo COMPLETO, "
    "ajustado con el 100% de los datos, no el de partición 80/20, porque DeLong exige "
    "comparación pareada sobre los mismos sujetos). Algoritmo estándar de DeLong et al. "
    "(1988; versión rápida de Sun & Xu, 2014), verificado contra sklearn (coincide "
    "exactamente en las 3 AUC)."
)
imagen("roc_comparacion_3_modelos_delong.png", width_in=5.4,
       caption="Figura 5.1. Curvas ROC de los 3 modelos sobre los mismos 256 puntos.")
tabla_desde_csv(f"{DIR_TAB}/delong_comparacion_pareada.csv", font_size=8.5)
interpretacion(
    "Las 3 comparaciones por pares son TODAS estadísticamente significativas (p<0.05): "
    "tanto WoE como Regresión Logística superan de forma significativa a AHP+Combinado, "
    "y Regresión Logística también supera de forma significativa a WoE. Esto confirma "
    "con evidencia estadística formal —no solo con el AUC puntual— que anclar los pesos "
    "a la evidencia del inventario (WoE) o ajustarlos por máxima verosimilitud "
    "(Regresión) mejora sustancialmente la capacidad de discriminación frente a "
    "asignarlos únicamente por criterio experto."
)
salto()

# ===========================================================================
# 6. Comparación final y conclusiones
# ===========================================================================
h("6. Comparación final y conclusiones", 1)
tabla_desde_csv(f"{DIR_TAB}/resumen_comparativo_AUC.csv", font_size=8.5)
interpretacion(
    "Ranking de menor a mayor capacidad de discriminación: AHP+Combinado "
    f"({auc_ahp:.3f}) < Regresión Logística validación ({auc_test:.3f}) < WoE "
    f"({auc_woe:.3f}) < Regresión Logística entrenamiento ({auc_train:.3f}). "
    "AHP+Combinado queda muy por debajo de WoE y de la Regresión Logística — apenas "
    "supera el azar y NO debería presentarse como el modelo final de susceptibilidad "
    "de este proyecto sin advertir esta limitación."
)

h("Recomendaciones", 2)
bullets([
    "El hallazgo más importante de este informe: al remover Frequency Ratio de "
    "AHP+Combinado (para ajustarse al ejemplo de referencia del curso), el AUC cayó de "
    "~0.79-0.82 (versión anterior, con FR) a ~0.54 (versión actual, sin FR) — un cambio "
    "metodológico con consecuencias grandes en el desempeño, que vale la pena explicar "
    "en la entrega si el profesor pregunta por qué el modelo heurístico puro rinde tan "
    "cerca del azar.",
    "Si el objetivo es un mapa de susceptibilidad con buena capacidad predictiva, usar "
    "WoE o Regresión Logística, no AHP+Combinado en su forma actual (sin FR).",
    "Para reportar la capacidad predictiva real de la regresión logística, usar el AUC "
    "de validación (0.729), no el de entrenamiento (0.898).",
    "La prueba de DeLong confirma con rigor estadístico (no solo con el AUC puntual) "
    "que WoE y Regresión superan al método heurístico puro — un argumento sólido para "
    "justificar la elección del modelo final ante el profesor.",
])
salto()

# ===========================================================================
# 7. Mapas de susceptibilidad — criterio SGC (tasa de éxito acumulada)
# ===========================================================================
h("7. Mapas de susceptibilidad — criterio del SGC (3 clases)", 1)
para(
    "El profesor pidió construir el mapa de susceptibilidad de cada método clasificando "
    "en 3 niveles según el criterio del SGC: Alta = 75% de los movimientos reales, "
    "Media = el siguiente 23% (acumulado 98%), Baja = el resto (acumulado 100%). El "
    "umbral de cada clase se obtiene de la curva ROC/de éxito: se ordenan los 128 "
    "movimientos reales por su score de mayor a menor, y se toma el score del movimiento "
    "que ocupa la posición del percentil correspondiente (p. ej. la posición 96 de 128 "
    "para el 75%) — ese score es el umbral que separa Alta de Media. El mismo criterio "
    "se aplicó a los 3 modelos usando su curva de éxito \"completa\" (100% de los datos, "
    "sin partición, igual que en la sección 5 de DeLong)."
)
resumen_area = pd.read_csv(f"{DIR_TAB}/resumen_area_3clases_todos_modelos.csv")
tabla_desde_csv(f"{DIR_TAB}/resumen_area_3clases_todos_modelos.csv", font_size=9)
interpretacion(
    "Por construcción, la clase \"Alta\" de los 3 mapas contiene exactamente el 75% de "
    "los movimientos reales (así se calibró el umbral) — la diferencia real entre "
    "modelos es CUÁNTA ÁREA de la cuenca necesita cada uno para capturar ese 75%: menos "
    "área = modelo más útil en la práctica, porque concentra el riesgo en una zona más "
    "pequeña y manejable. AHP+Combinado necesita el 69.5% de la cuenca en \"Alta\" — "
    "prácticamente inútil para priorizar zonas, coherente con su AUC apenas por encima "
    "del azar. WoE es el más concentrado (42.2% de área en \"Alta\"). Regresión "
    "Logística, pese a tener el AUC más alto (0.869), necesita más área que WoE (48.2%) "
    "para llegar al mismo 75% de movimientos — esto no es una contradicción: el AUC mide "
    "qué tan bien el modelo ORDENA los puntos en general, mientras que esta métrica mide "
    "qué tan CONCENTRADA queda la probabilidad en el mapa completo; con 3 de sus 5 "
    "variables categóricas (uso_actual, geología, geomorfología), la regresión asigna "
    "probabilidades casi idénticas a polígonos completos de esas categorías, generando "
    "bloques grandes y contiguos de \"Alta\" en vez de una zonificación más fina."
)
imagen("comparacion_3_mapas_susceptibilidad.png", width_in=6.3,
       caption="Figura 7.1. Comparación de los 3 mapas de susceptibilidad clasificados (SGC).")

for nombre_modelo, archivo_val, archivo_fig in [
    ("AHP + Combinado", "ahp_validacion_3clases.csv", "mapa_susceptibilidad_AHP_3clases.png"),
    ("WoE", "woe_validacion_3clases.csv", "mapa_susceptibilidad_WoE_3clases.png"),
    ("Regresión Logística", "regresion_validacion_3clases.csv", "mapa_susceptibilidad_Regresion_3clases.png"),
]:
    h(f"Mapa — {nombre_modelo}", 2)
    imagen(archivo_fig, width_in=4.6)
    tabla_desde_csv(f"{DIR_TAB}/{archivo_val}", font_size=9)

doc.save(OUT)
print(f"\nGuardado: {OUT}")
