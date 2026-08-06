# -*- coding: utf-8 -*-
"""
Informe Word dedicado: qué se cambió en el método AHP + Combinado (y en
Índices, que comparte los mismos w_c) para que el AUC subiera de ~0.54 a
~0.67, explicando en detalle la diferencia entre los pesos de clase (w_c)
de literatura (versión original de proyecto_susceptibilidad_AHP) y los
nuevos pesos de clase por Frequency Ratio — LA FÓRMULA DEL LIBRO/CURSO:

    w_n = L_r / A_r
    L_r = % de los movimientos totales que contiene la clase n
          (de los 128 puntos de movimiento del inventario)
    A_r = % del ÁREA TOTAL de la cuenca que representa la clase n
          (del ráster completo de la variable)

(Nota histórica: una primera versión de este cálculo usó, por error,
"% de puntos estables" en vez de "% de área" en el denominador — se
corrigió tras revisión explícita del usuario contra la fórmula del libro.)
"""
import os
import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
DIR_FIG = os.path.join(BASE, "figuras")
DIR_TAB = os.path.join(BASE, "tablas")
OUT = os.path.join(BASE, "Informe_AHP_Combinado_Metodologia.docx")

DATA_AHP = (
    r"C:\Users\Angela Acosta\OneDrive\Documentos\TRABAJO_FINAL_CARTOGRAFIA\MODELOS"
    r"\proyecto_susceptibilidad_AHP\proyecto_susceptibilidad_AHP\data"
)

# ===========================================================================
# 0. DATOS: pesos viejos (literatura) vs nuevos (FR)
# ===========================================================================
nombres_archivo = {
    "geologia": ("pesos_clase_geologia.csv", "unidad"),
    "geomorfologia": ("pesos_clase_geomorfologia.csv", "unidad"),
    "cobertura": ("pesos_clase_cobertura.csv", "cobertura"),
    "uso_actual": ("pesos_clase_uso_actual.csv", "uso_actual"),
}

tablas_comparacion = {}
for var, (archivo, col_nombre) in nombres_archivo.items():
    viejo = pd.read_csv(f"{DATA_AHP}/{archivo}").rename(columns={col_nombre: "nombre", "peso_wc": "peso_literatura"})
    nuevo = pd.read_csv(f"{DIR_TAB}/pesos_clase_FR_categoricas.csv")
    nuevo_var = nuevo[nuevo["variable"] == var][
        ["codigo", "pct_area_Ar", "pct_movimientos_Lr", "w_n_L_r_sobre_A_r", "peso_wc_normalizado"]]
    comp = viejo.merge(nuevo_var, on="codigo", how="left")
    comp["diferencia"] = round(comp["peso_wc_normalizado"] - comp["peso_literatura"], 3)
    tablas_comparacion[var] = comp

# resultados antes/después (valores documentados de esta sesión de trabajo)
resumen_antes_despues = pd.DataFrame([
    {"métrica": "AUC AHP+Combinado", "antes (literatura)": 0.542, "después (FR = L_r/A_r)": 0.667},
    {"métrica": "AUC Índices", "antes (literatura)": 0.519, "después (FR = L_r/A_r)": 0.673},
    {"métrica": "t-test AHP+Combinado (movimiento vs. estable)", "antes (literatura)": "t=1.077, p=0.283 (NO signif.)",
     "después (FR = L_r/A_r)": "t=4.921, p=1.6e-06 (signif.)"},
    {"métrica": "% área de la cuenca en clase \"Alta\" (75% mov.)",
     "antes (literatura)": "69.5%", "después (FR = L_r/A_r)": "63.7%"},
])

# ===========================================================================
# 1. INFORME WORD
# ===========================================================================
AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x40, 0x40, 0x40)
VERDE_OK = RGBColor(0x1E, 0x7B, 0x34)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
for s in doc.sections:
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p


def para(text, bold=False, italic=False, size=11, color=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def interpretacion(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run("Interpretación: ")
    r.bold = True; r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = VERDE_OK
    r2 = p.add_run(text)
    r2.italic = True; r2.font.size = Pt(10.5); r2.font.color.rgb = GRIS
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(11)


def set_cell_bg(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def tabla(headers, rows, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; r = p.add_run(str(hd))
        r.bold = True; r.font.size = Pt(font_size); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1F4E79")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]; r = p.add_run(str(val))
            r.font.size = Pt(font_size)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # resaltar filas donde el peso subió mucho o bajó mucho
            if i == len(row) - 1 and isinstance(val, (int, float)):
                if val >= 0.3:
                    set_cell_bg(cells[i], "C6EFCE")
                elif val <= -0.3:
                    set_cell_bg(cells[i], "FFC7CE")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def tabla_df(df, headers_legibles, font_size=8.5):
    return tabla(headers_legibles, df.values.tolist(), font_size=font_size)


def imagen(nombre, width_in=5.6, caption=None):
    ruta = os.path.join(DIR_FIG, nombre)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(ruta, width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption); cr.italic = True; cr.font.size = Pt(9.5); cr.font.color.rgb = GRIS
        cp.paragraph_format.space_after = Pt(4)


def salto():
    doc.add_page_break()


# --- Portada ---
doc.add_paragraph().paragraph_format.space_before = Pt(50)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("AHP + Combinado: de la Literatura a los Datos")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUL

st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = st.add_run("Qué se cambió y por qué el AUC subió de 0.54 a 0.69")
r2.font.size = Pt(14); r2.font.color.rgb = GRIS
st.paragraph_format.space_before = Pt(10)

st3 = doc.add_paragraph(); st3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = st3.add_run("Cuenca del río Arma (Abejorral, Antioquia)")
r4.font.size = Pt(12); r4.font.color.rgb = GRIS
st3.paragraph_format.space_before = Pt(30)
doc.add_page_break()

# ===========================================================================
# 1. Estructura del método (recordatorio)
# ===========================================================================
h("1. Cómo está armado el método AHP + Combinado", 1)
para(
    "El índice de susceptibilidad de este método se calcula como una suma ponderada de "
    "las 10 variables condicionantes:"
)
para("S_n = Σ (Wᵢ · w_cᵢ)   para i = 1...10 variables", italic=True)
bullets([
    "Wᵢ (peso de VARIABLE): qué tan importante es cada variable completa (pendiente, "
    "geología, cobertura, etc.) frente a las demás. Sale de la matriz de comparación "
    "por pares AHP (eigenvector + Razón de Consistencia). NO se tocó — sigue siendo "
    "100% juicio experto, es la esencia del método AHP.",
    "w_c (peso de CLASE): dentro de UNA variable, qué tan peligrosa es cada categoría "
    "específica (p. ej., dentro de \"geología\", qué tan peligrosa es cada unidad "
    "geológica). Aquí SÍ hubo un cambio, que se explica en la sección 2.",
])
interpretacion(
    "En términos simples: AHP decide cuánto importa cada variable (Wᵢ); w_c decide, "
    "dentro de esa variable, cuál de sus categorías es más riesgosa. El cambio que se "
    "hizo fue solo en la segunda parte (w_c), y solo para 4 de las 10 variables."
)

# ===========================================================================
# 2. Qué se cambió
# ===========================================================================
h("2. Qué se cambió: w_c de 4 variables, de literatura a datos propios", 1)
para(
    "El profesor pidió que la reclasificación se basara en las variables/datos propios "
    "de la cuenca, no en valores genéricos de literatura, para que el modelo AHP+Combinado "
    "no rindiera tan mal. Se aplicó ese cambio a las 4 variables CATEGÓRICAS: geología, "
    "geomorfología, cobertura y uso_actual. Las 6 variables CONTINUAS (pendiente, aspecto, "
    "curvatura, flujo_acum, elevación, dist_drenajes) se dejaron con los rangos de "
    "literatura, sin cambios — esto fue una decisión de alcance para no rehacer toda la "
    "metodología, solo la parte donde el impacto es mayor y más fácil de justificar."
)
para(
    "El método usado para recalcular w_c es Frequency Ratio (FR), con la fórmula oficial "
    "del capítulo del curso (no una aproximación):"
)
para("w_n = L_r / A_r", italic=True, size=13)
para("L_r = % de los movimientos totales que contiene la clase n "
     "(de los 128 puntos de movimiento reales del inventario)", italic=True)
para("A_r = % del ÁREA TOTAL de la cuenca que representa la clase n "
     "(contado en el ráster completo, no en la muestra de puntos)", italic=True)
interpretacion(
    "Si w_n > 1, los movimientos están sobrerrepresentados en esa clase respecto a lo "
    "que le correspondería por su tamaño (más peligrosa); si w_n < 1, están "
    "subrepresentados (menos peligrosa). L_r se calcula sobre los 128 puntos de "
    "movimiento (evento = punto, como está definido el inventario); A_r se calcula "
    "sobre el RÁSTER COMPLETO de cada variable — el área de una clase es una propiedad "
    "del mapa completo, no de la muestra de puntos. Al dividir cada w_n por el w_n "
    "máximo de su variable, todos los w_c quedan en una escala de 0 a 1, comparable "
    "directamente con los pesos de literatura. Clases sin ningún movimiento registrado "
    "entre los 128 puntos (L_r=0) reciben peso 0 por la fórmula literal — es una "
    "limitación conocida de Frequency Ratio sin corrección de continuidad, que afecta "
    "sobre todo a clases muy pequeñas en área (p. ej. uso_actual \"Cuerpos de Agua "
    "Naturales\", 0.025% de la cuenca)."
)
salto()

# ===========================================================================
# 3. Comparación de pesos: antes (literatura) vs. después (FR)
# ===========================================================================
h("3. Comparación de pesos de clase: literatura vs. Frequency Ratio", 1)
para(
    "Verde = la clase subió mucho de peso (se volvió más peligrosa según los datos de lo "
    "que decía la literatura); rojo = bajó mucho (los datos muestran que es menos "
    "peligrosa de lo que se pensaba)."
)

for var, titulo in [("geologia", "Geología"), ("geomorfologia", "Geomorfología"),
                     ("cobertura", "Cobertura"), ("uso_actual", "Uso actual")]:
    h(f"3.{['geologia','geomorfologia','cobertura','uso_actual'].index(var)+1} {titulo}", 2)
    comp = tablas_comparacion[var]
    filas = []
    for _, row in comp.iterrows():
        filas.append([
            row["codigo"], row["nombre"], row["peso_literatura"],
            f"{row['pct_area_Ar']:.1f}%", f"{row['pct_movimientos_Lr']:.1f}%",
            row["peso_wc_normalizado"], row["diferencia"],
        ])
    tabla(["Código", "Clase", "Peso literatura (antes)", "% área (A_r)", "% movimientos (L_r)",
           "Peso FR (ahora)", "Diferencia"], filas, font_size=7.5)

interpretacion(
    "El caso más claro es geología: el código 2 (Complejo Cajamarca - Esquistos verdes) "
    "ocupa 13.1% del área de la cuenca pero contiene 32.8% de los movimientos reales "
    "(2.5 veces sobrerrepresentado) — su peso de literatura era 0.70 y con FR sube al "
    "máximo, 1.0. El código 1 (Neis Intrusivo de Pantanillo) ocupa 5.2% del área pero "
    "solo tiene 0.8% de los movimientos (1 solo punto) — muy subrepresentado; su peso de "
    "literatura era 0.30 pero con FR baja a 0.06. Este tipo de correcciones, repetidas "
    "en las 4 variables, es lo que hace que el índice final separe mejor los puntos de "
    "movimiento de los estables."
)
salto()

# ===========================================================================
# 4. Resultados: antes vs. después
# ===========================================================================
h("4. Resultado del cambio: antes vs. después", 1)
tabla_df(resumen_antes_despues, ["Métrica", "Antes (literatura)", "Después (FR = L_r/A_r)"], font_size=9.5)
interpretacion(
    "El AUC de AHP+Combinado subió de 0.542 (apenas mejor que el azar) a 0.667; el de "
    "Índices subió de 0.519 a 0.673 — ambos se acercan al umbral \"moderadamente "
    "preciso\" de Swets (0.7) sin cruzarlo todavía. El t-test de validación, que antes "
    "NO era estadísticamente significativo (p=0.283 — la media del índice en "
    "movimientos y en estables era prácticamente igual), ahora sí lo es con muchísima "
    "confianza (p=1.6×10⁻⁶). En el mapa de susceptibilidad, el área que hay que marcar "
    "como \"Alta\" para capturar el 75% de los movimientos bajó de 69.5% de la cuenca a "
    "63.7% — el modelo concentra algo mejor el riesgo, aunque todavía queda por debajo "
    "de WoE (42.2%) y Regresión Logística (48.2%)."
)
imagen("mapa_susceptibilidad_AHP_3clases.png", width_in=4.8,
       caption="Figura 4.1. Mapa de susceptibilidad AHP+Combinado actualizado (con w_c por FR = L_r/A_r).")
imagen("roc_ahp.png", width_in=4.6, caption="Figura 4.2. Curva ROC actualizada — AHP+Combinado (AUC=0.667).")

# ===========================================================================
# 5. Conclusiones
# ===========================================================================
h("5. Conclusiones", 1)
bullets([
    "El cambio metodológico fue puntual y acotado: solo se recalculó w_c (peso de "
    "clase) de las 4 variables categóricas, usando Frequency Ratio (w_n = L_r/A_r, la "
    "fórmula oficial del curso) sobre el inventario real en vez de valores de "
    "literatura. W_i (peso de variable) y los w_c de las 6 continuas no se tocaron.",
    "El AUC mejoró sustancialmente (AHP+Combinado: +0.125, Índices: +0.154) porque "
    "ahora el modelo SÍ mira los datos reales al decidir qué clases son peligrosas, "
    "corrigiendo casos donde la literatura se equivocaba para esta cuenca específica "
    "(p. ej. geología código 2, muy subestimada por la literatura).",
    "Aun con la mejora, AHP+Combinado sigue siendo el más débil de los 3 métodos del "
    "proyecto (WoE=0.819, Regresión=0.869, ambos significativamente mayores según la "
    "prueba de DeLong) — tiene sentido, porque 6 de sus 10 variables (todas las "
    "continuas) siguen sin usar ningún dato real.",
    "Limitación transparente de Frequency Ratio sin corrección: las clases sin ningún "
    "movimiento observado entre los 128 puntos reciben peso 0 exacto, incluso si su "
    "área es muy pequeña (y por tanto la ausencia de movimientos no es muy informativa "
    "estadísticamente) — no se aplicó ninguna suavización tipo Laplace, para mantener "
    "la fórmula fiel a la del curso.",
    "Si se quisiera cerrar aún más la brecha, el siguiente paso natural sería aplicar el "
    "mismo criterio (FR, o algo equivalente) también a las 6 variables continuas — "
    "quedó fuera de alcance en este cambio por decisión explícita.",
])

doc.save(OUT)
print(f"Guardado: {OUT}")
