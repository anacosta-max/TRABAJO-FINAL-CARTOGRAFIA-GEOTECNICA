# -*- coding: utf-8 -*-
"""
Genera el informe Word de 13_Regresion_Logistica: metodología de datos
(GLM logístico) con selección de variables por pasos (forward, backward,
stepwise bidireccional) por criterio AIC, siguiendo el mismo formato de
informe ya usado en 08_Metodos_Conocimiento y 06_Variables_Condicionantes
(python-docx; docx-js/Node no disponible en esta máquina).
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, "figuras")
TAB = os.path.join(BASE, "tablas")
MAPAS = os.path.join(BASE, "mapas")
OUT = os.path.join(BASE, "Informe_Regresion_Logistica.docx")

AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x40, 0x40, 0x40)
VERDE_OK = RGBColor(0x1E, 0x7B, 0x34)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
for s in doc.sections:
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p


def para(text, bold=False, italic=False, size=11, color=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def interpretacion(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run("Interpretación: ")
    r.bold = True; r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = VERDE_OK
    r2 = p.add_run(text)
    r2.italic = True; r2.font.size = Pt(10.5); r2.font.color.rgb = GRIS
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(11)


def set_cell_bg(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def tabla(headers, rows, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; r = p.add_run(str(hd))
        r.bold = True; r.font.size = Pt(font_size); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1F4E79")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]; r = p.add_run(str(val))
            r.font.size = Pt(font_size)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def tabla_desde_csv(ruta_csv, font_size=8.5, decimales=3, max_filas=None, columnas=None):
    df = pd.read_csv(ruta_csv)
    if columnas:
        df = df[columnas]
    if max_filas:
        df = df.head(max_filas)
    df_fmt = df.copy()
    for col in df_fmt.columns:
        if pd.api.types.is_float_dtype(df_fmt[col]):
            df_fmt[col] = df_fmt[col].round(decimales)
    headers = [str(c) for c in df_fmt.columns]
    rows = df_fmt.values.tolist()
    return tabla(headers, rows, font_size=font_size)


def imagen(ruta_o_nombre, carpeta=FIG, width_in=6.2, caption=None):
    ruta = ruta_o_nombre if os.path.isabs(ruta_o_nombre) else os.path.join(carpeta, ruta_o_nombre)
    if not os.path.exists(ruta):
        para(f"[Figura no encontrada: {ruta}]", italic=True, color=RGBColor(0xC0, 0x00, 0x00))
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(ruta, width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption); cr.italic = True; cr.font.size = Pt(9.5); cr.font.color.rgb = GRIS
        cp.paragraph_format.space_after = Pt(4)


def salto():
    doc.add_page_break()


def toc():
    p = doc.add_paragraph(); run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = 'TOC \\o "1-2" \\h \\z \\u'
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Haga clic derecho y seleccione Actualizar campo para generar la tabla de contenido."
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, i, s, t, e):
        run._r.append(el)


NOMBRES_GEOLOGIA = {1: "Neis Intrusivo de Pantanillo", 2: "Complejo Cajamarca - Esquistos verdes",
                     3: "Neis Intrusivo Abejorral", 4: "Formación Abejorral",
                     6: "Complejo Cajamarca - Filitas/esquistos cuarzo-sericíticos y esquistos alumínicos"}
NOMBRES_GEOMORFOLOGIA = {1: "Espolón bajo de longitud larga", 2: "Espolón moderado de longitud larga",
                          3: "Espolón faceteado alto de longitud larga", 4: "Espolón alto de longitud media",
                          5: "Lomo denudado bajo de longitud larga", 6: "Espolón festoneado alto de longitud larga",
                          7: "Escarpe de línea de falla", 9: "Espolón alto de longitud larga",
                          10: "Espolón moderado de longitud media", 12: "Espolón bajo de longitud media"}
NOMBRES_USO_ACTUAL = {30202: "Cultivos transitorios semi-intensivos (CTS)",
                       30204: "Cultivos permanentes semi-intensivos (CPS)",
                       30206: "Pastoreo semi-intensivo (PSI)", 30207: "Pastoreo extensivo (PEX)",
                       30213: "Áreas para conservación/recuperación de la naturaleza (CRE)"}

# ===========================================================================
# PORTADA
# ===========================================================================
doc.add_paragraph().paragraph_format.space_before = Pt(60)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Métodos Basados en Datos")
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = AZUL

st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = st.add_run("Cartografía Geotécnica — Regresión Logística y Selección de Variables por Pasos")
r2.font.size = Pt(15); r2.font.color.rgb = GRIS
st.paragraph_format.space_before = Pt(12)

st2 = doc.add_paragraph(); st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = st2.add_run("Forward selection, backward elimination y stepwise bidireccional por AIC\n"
                  "Seguimiento del pseudo R² de McFadden al combinar variables")
r3.italic = True; r3.font.size = Pt(11.5); r3.font.color.rgb = GRIS
st2.paragraph_format.space_before = Pt(6)

st3 = doc.add_paragraph(); st3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = st3.add_run("Cuenca del río Arma (Abejorral, Antioquia)")
r4.font.size = Pt(12); r4.font.color.rgb = GRIS
st3.paragraph_format.space_before = Pt(40)

st4 = doc.add_paragraph(); st4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = st4.add_run("Julio de 2026")
r5.font.size = Pt(11); r5.font.color.rgb = GRIS

salto()
h("Tabla de contenido", 1)
toc()
salto()

# ===========================================================================
# 1. INTRODUCCIÓN
# ===========================================================================
h("1. Introducción y objetivo", 1)
para(
    "Este informe documenta la aplicación del capítulo \"Métodos basados en datos\" "
    "del libro de Cartografía Geotécnica, específicamente el modelo paramétrico de "
    "Regresión Logística (GLM binomial), con selección de variables por pasos "
    "(stepwise selection). A diferencia de los métodos heurísticos (AHP, Índices), "
    "aquí los pesos de las variables NO se fijan por juicio experto: se estiman "
    "estadísticamente a partir del inventario de 256 puntos (128 movimientos + 128 "
    "puntos NO), maximizando la verosimilitud de observar los datos reales."
)
para(
    "Siguiendo la indicación del profesor, se aplican las tres variantes clásicas de "
    "selección por pasos —forward selection, backward elimination y stepwise "
    "bidireccional— usando el AIC (Criterio de Información de Akaike) como criterio "
    "de decisión en cada paso, y registrando en cada uno el AIC y el pseudo R² de "
    "McFadden para ver cómo mejora (o no) el ajuste al combinar variables."
)
para(
    "Decisión metodológica: cada variable —numérica o categórica— entra o sale "
    "COMPLETA del modelo en cada paso (todas las categorías/dummies de una variable "
    "cualitativa se mueven juntas como un bloque). Esto evita la inestabilidad de "
    "tratar cada categoría individual como un predictor independiente con una muestra "
    "de apenas 256 observaciones.",
    bold=True,
)

# ===========================================================================
# 2. DATOS
# ===========================================================================
h("2. Datos", 1)
para(
    "Fuente: el mismo inventario discriminado y espacialmente estratificado usado en "
    "06_Variables_Condicionantes y en la validación de 08_Metodos_Conocimiento — 128 "
    "movimientos reales (puntos_modelo_balanceado_1.shp, sin modificar) y 128 puntos "
    "\"NO\" generados por disimilitud ambiental moderada con muestreo espacialmente "
    "estratificado (11_Inventario_Discriminado_Estratificado)."
)
tabla(
    ["Variable", "Tipo"],
    [
        ["pendiente", "Continua"], ["aspecto", "Continua (circular)"], ["curvatura", "Continua"],
        ["flujo_acum", "Continua"], ["elevación", "Continua"], ["dist_drenajes", "Continua"],
        ["cobertura", "Categórica (12 clases + Otras)"], ["uso_actual", "Categórica (6 clases + Otras)"],
        ["geología", "Categórica (5 clases + Otras)"], ["geomorfología", "Categórica (10 clases + Otras)"],
    ],
    font_size=9,
)
salto()

# ===========================================================================
# 3. PREPROCESAMIENTO
# ===========================================================================
h("3. Preprocesamiento: dos problemas de identificabilidad y su solución", 1)
para(
    "Antes de ajustar cualquier modelo, se probó el modelo \"completo\" con las 10 "
    "variables a la vez (el punto de partida clásico de backward elimination). Este "
    "modelo saturado presentó DOS problemas reales de identificabilidad, diagnosticados "
    "y corregidos de forma transparente:"
)

h("3.1 Separación (cuasi-)completa por categorías dispersas", 2)
para(
    "Varias categorías de las variables cualitativas tienen muy pocas observaciones y, "
    "además, un resultado homogéneo (todas \"movimiento\" o todas \"NO\"), por ejemplo "
    "cobertura=3.0 (n=1, 100% movimiento) o uso_actual código 30211 (n=2, 0% movimiento). "
    "Esto produce separación (cuasi-)completa: la verosimilitud no tiene máximo finito "
    "bien definido para esas categorías y los optimizadores numéricos o se quedan "
    "estancados o divergen a coeficientes artificialmente enormes."
)
para(
    "Solución aplicada: toda categoría con menos de 5 observaciones se agrupa en una "
    "clase \"Otras\"; si esa clase \"Otras\" resultante seguiría siendo muy pequeña "
    "(menos de 5 observaciones), esas observaciones se fusionan directamente con la "
    "categoría mayoritaria de la variable en su lugar."
)
tabla_desde_csv(os.path.join(TAB, "categorias_agrupadas.csv"), font_size=8.5)
interpretacion(
    "En total se agruparon 9 categorías dispersas entre las 4 variables cualitativas. "
    "El caso de geología ilustra la regla de fusión con la mayoritaria: la única "
    "categoría rara (código 5, n=1) se fusionó directamente con la categoría 3 "
    "(la más frecuente) en vez de crear su propia clase \"Otras\" de tamaño 1, que "
    "habría reproducido el mismo problema de separación que se buscaba corregir."
)

h("3.2 Colinealidad casi perfecta entre cobertura y uso_actual", 2)
para(
    "Incluso después de agrupar las categorías dispersas, el modelo con las 10 "
    "variables seguía sin converger de forma sana. La causa: \"cobertura\" (cobertura "
    "vegetal) y \"uso_actual\" (uso actual del suelo) están casi perfectamente "
    "anidadas en esta muestra — cada clase de cobertura corresponde casi 1 a 1 con una "
    "clase de uso actual (tabla siguiente). Esto deja la matriz de diseño con rango "
    "deficiente cuando ambas entran juntas: el modelo no es identificable."
)
tabla_desde_csv(os.path.join(TAB, "diagnostico_cobertura_uso_actual.csv"), font_size=7.5)
interpretacion(
    "La correspondencia es casi biunívoca: cobertura=1 ⇄ uso_actual=30204 (74 de 74 "
    "casos), cobertura=2 ⇄ uso_actual=30206 (79 de 79), cobertura=4 ⇄ uso_actual=30207 "
    "(40 de 40), cobertura=8 ⇄ uso_actual=30213 (27 de 27). En la práctica, estas dos "
    "variables capturan la misma información en esta cuenca (la cobertura vegetal "
    "observada es, casi siempre, consecuencia directa del uso del suelo declarado). "
    "Solución aplicada: en vez de excluir una de las dos a mano, se agregó una "
    "verificación automática de rango de la matriz de diseño antes de cada ajuste; "
    "cualquier combinación de variables no identificable se descarta igual que un "
    "modelo que no converge. Esto hace que la selección por pasos evite "
    "automáticamente incluir ambas variables a la vez, y el AIC decide cuál de las "
    "dos aporta más (ver sección 4)."
)
salto()

# ===========================================================================
# 4. MODELO DE REFERENCIA (9 VARIABLES)
# ===========================================================================
h("4. Modelo de referencia y punto de partida real de backward elimination", 1)
para(
    "Como el modelo con las 10 variables no es identificable, se calcularon los dos "
    "modelos posibles de 9 variables (excluyendo cobertura o excluyendo uso_actual) "
    "como referencia — el mejor de los dos es, en la práctica, el verdadero punto de "
    "partida de backward elimination (sección 6)."
)
tabla_desde_csv(os.path.join(TAB, "modelos_referencia_9_variables.csv"), font_size=8)
interpretacion(
    "El modelo sin \"cobertura\" (AIC=274.87) supera ligeramente al modelo sin "
    "\"uso_actual\" (AIC=282.90) — es decir, entre las dos variables casi redundantes, "
    "\"uso_actual\" aporta más información sobre la ocurrencia de movimientos que "
    "\"cobertura\" en esta cuenca. Ambos modelos de 9 variables mejoran sustancialmente "
    "el AIC del modelo nulo (356.89), con un pseudo R² de McFadden de ≈0.366."
)
salto()

# ===========================================================================
# 5. FORWARD SELECTION
# ===========================================================================
h("5. Forward selection (selección hacia adelante)", 1)
para(
    "Se parte del modelo nulo (solo intercepto) y en cada paso se agrega, de entre las "
    "variables aún no incluidas, la que produce el modelo con menor AIC. Se detiene "
    "cuando ninguna variable candidata mejora el AIC del modelo actual."
)
tabla_desde_csv(os.path.join(TAB, "forward_selection_historial.csv"), font_size=7.5,
                 columnas=["paso", "accion", "variable", "n_variables", "AIC",
                           "pseudo_R2_McFadden", "delta_AIC", "LR_p_valor"])
interpretacion(
    "El AIC baja de 356.89 (nulo) a 268.93 en 5 pasos, y el pseudo R² sube de 0 a "
    "0.3606. El mayor salto ocurre al agregar uso_actual (ΔAIC=-41.87, R² pasa a "
    "0.14) y se mantiene fuerte con geología (ΔAIC=-24.21, R²=0.23) y geomorfología "
    "(ΔAIC=-17.38, R²=0.34) — las tres variables categóricas dominan la capacidad "
    "predictiva del modelo. curvatura aporta una mejora moderada pero significativa "
    "(ΔAIC=-4.22, p=0.013). El último paso (dist_drenajes) mejora el AIC solo "
    "marginalmente (ΔAIC=-0.28) y su prueba de razón de verosimilitud ya no es "
    "significativa al 5% (p=0.131) — el AIC lo acepta porque reduce el criterio en "
    "términos absolutos, aunque sea por muy poco, pero esto anticipa que "
    "dist_drenajes probablemente no será significativo en el modelo final (sección 8). "
    "Nótese que \"pendiente\" —la variable de mayor peso en la jerarquía de literatura "
    "usada en 08_Metodos_Conocimiento— nunca entra al modelo: una vez que uso_actual, "
    "geología y geomorfología ya están incluidas, pendiente no aporta suficiente "
    "información adicional (ver también el hallazgo similar de 06_Variables_Condicionantes)."
)
salto()

# ===========================================================================
# 6. BACKWARD ELIMINATION
# ===========================================================================
h("6. Backward elimination (eliminación hacia atrás)", 1)
para(
    "Se parte del mejor modelo de 9 variables (sección 4, ya que el de 10 no es "
    "identificable) y en cada paso se quita la variable cuya eliminación más reduce "
    "el AIC, hasta que quitar cualquier otra variable ya no mejora el criterio."
)
tabla_desde_csv(os.path.join(TAB, "backward_elimination_historial.csv"), font_size=7.5,
                 columnas=["paso", "accion", "variable", "n_variables", "AIC",
                           "pseudo_R2_McFadden", "delta_AIC", "LR_p_valor"])
interpretacion(
    "Partiendo de las 9 variables (sin cobertura, AIC=274.87), se eliminan "
    "sucesivamente flujo_acum, pendiente, elevación y aspecto — cada una de estas "
    "eliminaciones en realidad MEJORA el AIC (todos los ΔAIC son negativos), lo que "
    "confirma que ninguna de esas cuatro variables aportaba información útil una vez "
    "que las demás ya estaban en el modelo (todas sus pruebas LR tienen p > 0.19, muy "
    "lejos de ser significativas). El pseudo R² prácticamente no cambia al quitarlas "
    "(de 0.3664 a 0.3606) — la pérdida de ajuste es mínima frente a la simplificación "
    "del modelo, exactamente el objetivo de backward elimination."
)
salto()

# ===========================================================================
# 7. STEPWISE BIDIRECCIONAL
# ===========================================================================
h("7. Stepwise bidireccional", 1)
para(
    "Combina forward y backward: en cada paso se evalúan tanto las variables que "
    "podrían agregarse como las que ya están y podrían quitarse, y se toma la acción "
    "que más reduce el AIC. Se detiene cuando ninguna acción (agregar o quitar) "
    "mejora el AIC actual."
)
tabla_desde_csv(os.path.join(TAB, "stepwise_bidireccional_historial.csv"), font_size=7.5,
                 columnas=["paso", "accion", "variable", "n_variables", "AIC",
                           "pseudo_R2_McFadden", "delta_AIC", "LR_p_valor"])
interpretacion(
    "El recorrido de stepwise bidireccional es idéntico al de forward selection "
    "(agrega uso_actual, geología, geomorfología, curvatura y dist_drenajes, en ese "
    "orden, sin retirar ninguna variable en el camino) — en este dataset, una vez que "
    "una variable entra por tener el mayor beneficio marginal de AIC, ninguna "
    "combinación posterior la vuelve prescindible, por lo que no hay necesidad de "
    "retroceder."
)
salto()

# ===========================================================================
# 8. COMPARACIÓN Y MODELO FINAL
# ===========================================================================
h("8. Comparación de los 3 métodos y modelo final", 1)
tabla_desde_csv(os.path.join(TAB, "comparacion_3_metodos.csv"), font_size=8)
interpretacion(
    "Los tres métodos de selección por pasos —forward selection, backward elimination "
    "y stepwise bidireccional— convergen EXACTAMENTE al mismo subconjunto de 5 "
    "variables (uso_actual, geología, geomorfología, curvatura, dist_drenajes), con el "
    "mismo AIC final (268.93) y el mismo pseudo R² (0.3606), a pesar de partir de "
    "puntos de arranque opuestos (vacío vs. casi-completo) y de recorrer caminos "
    "distintos. Esta coincidencia es una evidencia sólida de que el modelo de 5 "
    "variables es un óptimo robusto por AIC en este dataset, y no un artefacto del "
    "orden de búsqueda de un solo algoritmo."
)

h("8.1 Coeficientes del modelo final y odds ratios", 2)
para(
    "Fórmula: inventario ~ C(uso_actual) + C(geología) + C(geomorfología) + curvatura "
    "+ dist_drenajes. Para las variables categóricas, cada coeficiente representa el "
    "efecto de esa categoría FRENTE A la categoría de referencia (la primera en orden "
    "numérico tras agrupar categorías raras): geología código 1 (Neis Intrusivo de "
    "Pantanillo), geomorfología código 1 (Espolón bajo de longitud larga) y uso_actual "
    "código \"Otras\" (fusión de los usos forestales productor/protector, minoritarios "
    "en la muestra)."
)
tabla_desde_csv(os.path.join(TAB, "modelo_final_coeficientes.csv"), font_size=6.8, decimales=3)
interpretacion(
    "De los 21 parámetros del modelo final, 8 son estadísticamente significativos al "
    "5% (columna significativo_p<0.05 = True), además del intercepto marginal (p=0.083). "
    "Los efectos más fuertes: uso_actual código 30213 — \"Áreas para conservación/"
    "recuperación de la naturaleza\" — multiplica las probabilidades de movimiento por "
    "22 respecto a la categoría de referencia (OR=22.04, p<0.001, IC95%: 4.1-117); "
    "geología código 2 — \"Complejo Cajamarca - Esquistos verdes\" — por 26.4 (p=0.023); "
    "y geología código 6 — \"Complejo Cajamarca - Filitas/esquistos\" — por 24.3 "
    "(p=0.034). En geomorfología, dos categorías reducen significativamente el riesgo "
    "frente a la referencia: código 4 (\"Espolón alto de longitud media\", OR=0.094, "
    "p=0.014) y código 6 (\"Espolón festoneado alto de longitud larga\", OR=0.098, "
    "p=0.003), mientras código 5 (\"Lomo denudado bajo de longitud larga\") lo "
    "aumenta (OR=22.6, p=0.046). curvatura es significativa (OR=1.19 por unidad, "
    "p=0.015): a mayor curvatura (terreno más cóncavo), mayor probabilidad de "
    "movimiento, consistente con la concentración de agua e infiltración en zonas "
    "cóncavas. En cambio dist_drenajes, pese a haber sido seleccionada por los tres "
    "métodos por reducir el AIC, NO alcanza significancia individual al 5% (p=0.135) "
    "en el modelo conjunto — su aporte marginal al AIC (ΔAIC=-0.28, ver sección 5) ya "
    "anticipaba un efecto débil; se mantiene en el modelo porque el criterio de "
    "selección fue AIC y no un umbral de p-valor, pero esto se señala aquí como una "
    "limitación a tener en cuenta."
)

h("8.2 Desempeño: curva ROC y AUC", 2)
imagen("roc_modelo_final.png", width_in=4.8, caption="Figura 8.1. Curva ROC del modelo final.")
interpretacion(
    "El área bajo la curva ROC (AUC) del modelo final, evaluada sobre la misma "
    "muestra de ajuste (n=256, sin partición externa por el tamaño limitado de la "
    "muestra), es de 0.869 — una discriminación buena a muy buena entre puntos de "
    "movimiento y puntos NO según los umbrales convencionales (AUC>0.8 = bueno, "
    "AUC>0.9 = excelente). Es sustancialmente mejor que la validación de los métodos "
    "heurísticos puros de 08_Metodos_Conocimiento (26.6% y 25.8% de movimientos en "
    "zonas Alta+Muy alta para AHP e Índices), y se acerca al desempeño del método "
    "combinado con Frequency Ratio (48.4%), reflejando la ventaja de estimar los "
    "pesos directamente de la evidencia estadística del inventario en vez de fijarlos "
    "por juicio experto."
)
salto()

# ===========================================================================
# 9. MAPA DE SUSCEPTIBILIDAD
# ===========================================================================
h("9. Mapa de susceptibilidad", 1)
para(
    "Se aplica el modelo final a los rasters de las 5 variables seleccionadas en todo "
    "el dominio de la cuenca, reclasificando las categorías cualitativas con las "
    "mismas reglas de agrupación usadas al entrenar el modelo (sección 3.1). Las "
    "celdas cuya categoría no existe de ninguna forma en el modelo entrenado quedan "
    "como NoData — el modelo no puede extrapolar a categorías que nunca observó."
)
imagen("mapa_susceptibilidad_regresion_logistica.png", width_in=5.3,
       caption="Figura 9.1. Mapa de susceptibilidad — Regresión logística (5 clases por quintiles de probabilidad).")
tabla_desde_csv(os.path.join(TAB, "distribucion_area_clases_susceptibilidad.csv"), font_size=9)
interpretacion(
    "Las clases se definieron por quintiles de probabilidad predicha (20% del área en "
    "cada una), por lo que la distribución de área es, por construcción, uniforme. Lo "
    "relevante es dónde caen los movimientos REALES dentro de esas clases (tabla "
    "siguiente), no la distribución de área en sí."
)
tabla_desde_csv(os.path.join(TAB, "validacion_movimientos_en_clases.csv"), font_size=9)
interpretacion(
    "El 71.1% de los 128 movimientos reales (48.4% en \"Muy alta\" + 22.7% en "
    "\"Alta\") caen en las dos clases de mayor susceptibilidad, que en conjunto cubren "
    "solo el 40% del área de la cuenca — una concentración muy superior a la de los "
    "métodos heurísticos puros (26.6% AHP, 25.8% Índices) y comparable a la del "
    "método combinado con Frequency Ratio (48.4% solo en Alta+Muy alta) de "
    "08_Metodos_Conocimiento. Esto confirma, de forma independiente y con un método "
    "estadístico distinto (regresión logística multivariada vs. combinación lineal "
    "ponderada), que anclar los pesos a evidencia estadística local mejora "
    "sustancialmente la capacidad de discriminación frente al criterio puramente "
    "experto."
)
salto()

# ===========================================================================
# 10. CONCLUSIONES Y LIMITACIONES
# ===========================================================================
h("10. Conclusiones y limitaciones", 1)
bullets([
    "El modelo con las 10 variables NO es identificable en este dataset: cobertura y "
    "uso_actual están casi perfectamente anidadas (colinealidad severa), y varias "
    "categorías cualitativas son demasiado escasas para estimarse de forma estable. "
    "Ambos problemas se diagnosticaron con evidencia (tablas de contingencia, conteos "
    "por categoría) y se corrigieron de forma sistemática (agrupación de categorías "
    "raras, verificación de rango de la matriz de diseño) en vez de ocultarse.",
    "Los tres métodos de selección por pasos (forward, backward, stepwise "
    "bidireccional) convergen EXACTAMENTE al mismo modelo de 5 variables: uso_actual, "
    "geología, geomorfología, curvatura y dist_drenajes — evidencia de un óptimo "
    "robusto por AIC, no dependiente del algoritmo de búsqueda.",
    "El AIC baja de 356.89 (modelo nulo) a 268.93, y el pseudo R² de McFadden sube de "
    "0 a 0.3606 — una mejora sustancial de ajuste explicada casi en su totalidad por "
    "las tres variables categóricas (uso_actual, geología, geomorfología).",
    "pendiente, el factor de mayor peso en la jerarquía de literatura usada en "
    "08_Metodos_Conocimiento, nunca entra al modelo estadístico: una vez que "
    "uso_actual, geología y geomorfología están incluidas, no aporta información "
    "adicional relevante en esta cuenca — coincide con el hallazgo ya documentado en "
    "06_Variables_Condicionantes.",
    "dist_drenajes fue seleccionada por los tres métodos (mejora el AIC), pero no "
    "alcanza significancia individual al 5% (p=0.135) en el modelo conjunto — una "
    "limitación explícita a considerar si se buscara un modelo más parsimonioso "
    "basado en significancia estadística en vez de AIC.",
    "El modelo final tiene un AUC de 0.869 y concentra el 71.1% de los movimientos "
    "reales en las clases Alta+Muy alta de susceptibilidad (que cubren solo 40% del "
    "área) — un desempeño de validación muy superior a los métodos heurísticos puros "
    "(AHP 26.6%, Índices 25.8%) y comparable al método combinado con Frequency Ratio "
    "(48.4%).",
    "Limitación general: el AUC y la validación se calcularon sobre la misma muestra "
    "usada para ajustar el modelo (n=256), sin partición externa de entrenamiento/"
    "prueba, dado el tamaño limitado del inventario — el desempeño reportado puede "
    "estar optimistamente sesgado y debería confirmarse con validación cruzada o un "
    "conjunto de prueba independiente si se dispusiera de más puntos de movimiento.",
])

h("Recomendaciones", 2)
bullets([
    "Si se requiere un modelo más simple y fácilmente interpretable, considerar "
    "retirar dist_drenajes (la única variable seleccionada por AIC pero no "
    "significativa al 5%) y evaluar el modelo de 4 variables resultante.",
    "Dado que cobertura y uso_actual son casi redundantes en esta cuenca, no es "
    "necesario levantar ambas capas en trabajos futuros de la misma zona: uso_actual "
    "demostró aportar más información al modelo.",
    "Complementar este análisis con validación cruzada (k-fold) o partición "
    "entrenamiento/prueba si en el futuro se amplía el inventario de movimientos, "
    "para obtener una estimación del AUC menos optimista que la reportada aquí.",
])

doc.save(OUT)
print(f"Guardado: {OUT}")
